"""Add a second source document and a real cross-document oracle question."""

from __future__ import annotations

from pathlib import Path

from memory_data_service.schemas import FactRecord, QuestionRecord


def add_cross_document_case(
    *,
    dataset_id: str,
    dataset_path: Path,
    facts: list[FactRecord],
    questions: list[QuestionRecord],
    language: str = "en",
) -> Path | None:
    """Materialize a companion DOCX and a question that requires both sources."""
    if not facts:
        return None
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required for cross-document cases") from exc
    # A cross-document question must refer to an answer-bearing value, not to
    # an arbitrary first record.  Governance facts are narrative ownership
    # statements; calling them a "standard value" produces a semantically
    # inconsistent question, especially in the Chinese corpus.
    source = next(
        (
            fact
            for fact in facts
            if fact.fact_type in {"direct_numeric", "table_cell", "equation"}
            and fact.answer.strip()
        ),
        facts[0],
    )
    companion_fact_id = "FACT-CROSS-00001"
    is_chinese = language == "zh"
    companion_answer = "已启用" if is_chinese else "enabled"
    companion = dataset_path / f"{dataset_id}-companion.docx"
    document = Document()
    document.core_properties.title = (
        "LightRAG 合成辅助证据" if is_chinese else "LightRAG Synthetic Companion Evidence"
    )
    section = "辅助证据登记" if is_chinese else "Companion Evidence Register"
    expected_text = (
        f"{companion_fact_id}：辅助文档的核验状态为{companion_answer}。"
        if is_chinese
        else f"{companion_fact_id}: The companion verification state is {companion_answer}."
    )
    document.add_heading(section, 0)
    document.add_paragraph(expected_text)
    document.save(companion)
    facts.append(
        FactRecord(
            fact_id=companion_fact_id,
            fact_type="cross_document",
            answer=companion_answer,
            expected_text=expected_text,
            section=section,
            page=1,
            object_type="text",
            object_id_hint="companion.docx",
        )
    )
    questions.append(
        QuestionRecord(
            id="Q-CROSS-DOCUMENT-00001",
            question=(
                f"结合主文档与辅助文档，{source.fact_id}的核定结果和辅助文档的核验状态分别是什么？"
                if is_chinese
                else (
                    "Using the primary document and the companion document, what are the "
                    f"authoritative result for {source.fact_id} and the companion verification state?"
                )
            ),
            answer=f"{source.answer}; {companion_answer}",
            question_type="cross_document",
            evidence_fact_ids=[source.fact_id, companion_fact_id],
        )
    )
    return companion
