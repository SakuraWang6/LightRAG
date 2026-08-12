"""Deterministic Simplified-Chinese document generator for memory evaluation.

The generator produces thesis-grade synthetic Chinese documents: every page
carries a workstream-specific spec with concrete parameter values and units,
fluent narrative paragraphs (no verbatim per-page repetition), richer
parameter tables, per-page figures and parameterised equations.  Facts and
questions stay in the regular dataset contract so the evaluation framework
consumes them unchanged.
"""

from __future__ import annotations

import hashlib
import logging
import random
import time
from pathlib import Path

from memory_data_service.cross_document import add_cross_document_case
from memory_data_service.generators.docx_generator import (
    _convert_pdf,
    _file_record,
    _skipped,
)
from memory_data_service.provenance import (
    annotate_question_scenarios,
    build_provenance,
    resolve_scenario_quotas,
)
from memory_data_service.resource_guard import (
    GenerationResourceMonitor,
    enforce_generation_limits,
    estimate_generation_resources,
)
from memory_data_service.schemas import (
    DatasetCreateRequest,
    DatasetManifest,
    DocumentObject,
    FactRecord,
    GeneratedFile,
    ObjectRelation,
    OraclePayload,
    QuestionRecord,
)
from memory_data_service.storage import DEFAULT_GENERATED_ROOT, ensure_root, write_json

logger = logging.getLogger(__name__)

_DEFAULT_ENGLISH_TITLE = "LightRAG Synthetic Rich Memory Document"
_DEFAULT_CHINESE_TITLE = "LightRAG 合成中文记忆测评文档"

# Twelve workstreams so the narrative changes across a full cycle; pages that
# are far apart still share a theme but carry page-specific details.
_WORKSTREAMS = (
    (
        "售后知识服务与自动答复",
        "面向高频售后咨询构建可追溯的知识服务，范围内问题自动答复，涉及退款、隐私或高风险承诺的请求转入人工。",
    ),
    (
        "数据治理与来源准入",
        "知识条目须附责任人、有效期与审批记录，新增资料先经重复检测与敏感字段检查，再进入待发布版本。",
    ),
    (
        "发布节奏与灰度验证",
        "发布批次先在灰度环境核验召回、引用与拒答表现，未达门槛的批次只能修订，不得跳过评审直接上线。",
    ),
    (
        "运行监控与事件响应",
        "运行期持续监控答案依据、人工转接率与版本漂移，事件关闭必须记录根因、补救措施与后续验证窗口。",
    ),
    (
        "检索质量与相关性评估",
        "检索结果按证据命中率与排序质量打分，低分模板触发重新切片与索引重建，避免陈旧分块持续影响回答。",
    ),
    (
        "安全合规与隐私保护",
        "涉及隐私与退款的操作受双重签核约束，日志留存满足合规期限，敏感字段在检索与展示链路中脱敏。",
    ),
    (
        "模型版本与回滚管理",
        "模型版本与提示词基线绑定发布，线上表现劣化时按预案回滚，回滚前后保留可比对的评测快照。",
    ),
    (
        "容量规划与性能预算",
        "按峰值负载预留吞吐与延迟预算，扩容触发阈值由历史百分位校准，超预算请求进入降级队列。",
    ),
    (
        "多语言与跨区域服务",
        "同一知识库经翻译与本地化后服务多区域，跨区域引用以来源语言为准，禁止在转译层改写数值。",
    ),
    (
        "审计追溯与证据链",
        "每次回答保留引用来源与上下文快照，审计可按问题回溯到原始文档、分块与模型版本。",
    ),
    (
        "知识时效与过期清理",
        "知识条目带有效期与复核周期，过期内容先下线再修订，修订期间旧版本仅保留只读副本。",
    ),
    (
        "成本控制与资源配额",
        "查询与存储成本按配额核算，超额部分进入审批队列，常规场景优先复用缓存结果。",
    ),
)

# Per-page parameter profiles rotate so the document is not dominated by one
# repeated "standard calibration limit" sentence.
_PARAMETER_PROFILES = (
    ("响应时延预算", "ms", lambda p, r: f"{40 + p * 2 + r.randint(0, 6)}"),
    ("吞吐配额", "次/秒", lambda p, r: f"{100 + p * 5 + r.randint(0, 9)}"),
    ("错误率上限", "%", lambda p, r: f"{0.3 + p * 0.05 + r.randint(0, 4) * 0.05:.2f}"),
    ("转人工阈值", "分", lambda p, r: f"{25 + p + r.randint(0, 4)}"),
    ("校准窗口", "小时", lambda p, r: f"{4 + (p % 4) * 3 + r.randint(0, 2)}"),
    ("标准标定上限", "QMU", lambda p, r: f"{1000 + p * 7 + r.randint(0, 6)}"),
    ("冗余系数", "", lambda p, r: f"{1.1 + (p % 5) * 0.1 + r.randint(0, 2) * 0.05:.2f}"),
    ("保留期限", "天", lambda p, r: f"{30 + p * 3 + r.randint(0, 5)}"),
)

_SPEC_TEMPLATES = (
    (
        "{fact_id}：第 {page} 页{workstream}的{param}为 {value} {unit}。"
        "该值由上一发布周期的回归结果校准，纳入版本审计后作为本页基线。"
    ),
    (
        "{fact_id}：第 {page} 页{workstream}将{param}核定在 {value} {unit}。"
        "基线记录已归档，任何偏离该基线的数值都须重新评审。"
    ),
    (
        "{fact_id}：第 {page} 页{workstream}的{param}最终确定为 {value} {unit}，"
        "并在验收报告中标注为生效版本。"
    ),
    (
        "{fact_id}：第 {page} 页{workstream}的{param}为 {value} {unit}，"
        "该数值与上期基线一致，保留完整追溯链。"
    ),
)

_NARRATIVE_TEMPLATES = (
    (
        "围绕{workstream}，本页把{param}落实到 {value} {unit}，并同步更新了运行基线。"
        "配套的评审记录与来源编号一并归档，后续页面引用时直接回查本页。"
    ),
    (
        "{workstream}的{param}按 {value} {unit} 执行后，线上指标进入观察窗口。"
        "观察期内任何偏离都会触发告警并生成事件工单。"
    ),
    (
        "本页把{workstream}的{param}设定为 {value} {unit}。"
        "该决定与第 {prev_page} 页的基线一起构成连续审计链。"
    ),
    (
        "{workstream}将{param}维持在 {value} {unit}，配合容量预留与降级队列，"
        "确保高峰时段仍满足服务目标。"
    ),
)

_CONTROL_TEMPLATES = (
    "执行控制：{param}（{value} {unit}）先核验标准来源再记录变更，不带事实编号的数值不能作为最终结论。",
    "执行控制：{param}的调整必须附带来源编号与评审记录，未经验证的数值一律视为草稿。",
    "执行控制：{param}变更须在灰度环境完成回归验证，并保留前后两版的可比对快照。",
    "执行控制：涉及{param}的引用须回到原始事实条目，转述或缩写不得改变量纲。",
    "执行控制：超过{param}预算或阈值的请求进入降级与人工复核队列，不直接返回缓存结果。",
    "执行控制：审计日志与版本基线同步归档，回滚后仍可重建{workstream}在{param}为 {value} {unit} 时的上下文快照。",
)

_DISTRACTOR_TEMPLATES = (
    "草案阶段的近似值为 {near} {unit}，未通过评审，不得作为结论依据。",
    "早期草稿曾记录 {near} {unit}，因口径不一致被驳回，正式版本以本页核定值为准。",
)

_TABLE_THEMES = (
    (
        "时延阈值",
        "在表 {page} 中，标准行标记对应的最大值是多少？",
        "ms",
        lambda p, r: f"{p * 3 + 0.5}",
    ),
    (
        "吞吐配额",
        "在表 {page} 中，标准行标记对应的配额是多少？",
        "次/秒",
        lambda p, r: f"{p * 40 + 80}",
    ),
    (
        "错误率预算",
        "在表 {page} 中，标准行标记对应的预算上限是多少？",
        "%",
        lambda p, r: f"{p * 0.2 + 0.5:.1f}",
    ),
)

_FIGURE_STATES = (
    "已核验的控制流状态",
    "灰度到生产的发布门禁状态",
    "跨区域流量调度状态",
    "知识版本回滚状态",
)

_EQUATION_POOL = (
    ("响应时延", "L = L_r + L_k + L_g", "L_r 为检索耗时，L_k 为关键词与重排序耗时，L_g 为生成耗时"),
    ("吞吐配额", "T = W ÷ (t_s + t_o)", "W 为并行工作数，t_s 为单请求服务耗时，t_o 为固定开销"),
    ("服务可用性", "A = M ÷ (M + R)", "M 为平均无故障时间，R 为平均修复时间"),
    ("证据召回率", "R = H ÷ (H + M)", "H 为命中证据数，M 为未命中证据数"),
    ("单位服务成本", "C = Q × C_q + S × C_s", "Q 为查询量，C_q 为单次查询成本，S 为存储量，C_s 为单单位存储成本"),
)


def generate_chinese_dataset(
    request: DatasetCreateRequest,
    *,
    root: Path = DEFAULT_GENERATED_ROOT,
) -> DatasetManifest:
    """Generate a Chinese corpus while preserving the regular dataset contract."""
    started = time.perf_counter()
    root = ensure_root(root)
    pages = request.resolved_pages()
    enforce_generation_limits(request, pages)
    resource_estimate = estimate_generation_resources(request, pages)
    dataset_id = request.dataset_id or _stable_dataset_id(request, pages)
    dataset_path = root / dataset_id
    dataset_path.mkdir(parents=True, exist_ok=True)
    effective_request = request.model_copy(
        update={
            "title": (
                _DEFAULT_CHINESE_TITLE
                if request.title == _DEFAULT_ENGLISH_TITLE
                else request.title
            )
        }
    )
    docx_path = dataset_path / f"{dataset_id}.docx"

    with GenerationResourceMonitor() as resource_monitor:
        facts, questions, objects, relations = _write_docx(
            effective_request, docx_path, pages
        )
        companion_docx = (
            add_cross_document_case(
                dataset_id=dataset_id,
                dataset_path=dataset_path,
                facts=facts,
                questions=questions,
                language="zh",
            )
            if "docx" in effective_request.formats
            else None
        )
        scenario_counts = annotate_question_scenarios(questions)
        scenario_quotas = resolve_scenario_quotas(
            requested=effective_request.scenario_quotas, observed=scenario_counts
        )
        pdf_record = (
            _convert_pdf(docx_path, dataset_path)
            if "pdf" in effective_request.formats
            else _skipped("pdf")
        )
        oracle = OraclePayload(
            dataset_id=dataset_id,
            language="zh",
            facts=facts,
            questions=questions,
            objects=objects,
            relations=relations,
        )
        _write_artifacts(
            dataset_path=dataset_path,
            dataset_id=dataset_id,
            facts=facts,
            questions=questions,
            objects=objects,
            relations=relations,
            oracle=oracle,
        )

    files: list[GeneratedFile] = [
        _file_record(docx_path, "docx", role="source_document")
        if "docx" in effective_request.formats
        else _skipped("docx")
    ]
    files.append(pdf_record)
    if companion_docx is not None:
        files.append(_file_record(companion_docx, "docx", role="source_document"))
    for name in ("facts.json", "questions.json", "objects.json", "relations.json", "oracle.json"):
        files.append(_file_record(dataset_path / name, "json"))
    for asset_path in sorted(dataset_path.glob("*.png")):
        files.append(_file_record(asset_path, "png"))

    fingerprint, provenance = build_provenance(
        request=effective_request,
        pages=pages,
        generator="chinese_docx_generator",
        template_version=f"zh-{effective_request.profile}-docx-v3",
        source_file=Path(__file__),
    )
    manifest = DatasetManifest(
        dataset_id=dataset_id,
        tier=effective_request.tier,
        pages=pages,
        profile=effective_request.profile,
        language="zh",
        formats=effective_request.formats,
        modalities=effective_request.modalities,
        title=effective_request.title,
        display_name=effective_request.display_name,
        split=effective_request.split,
        scenario_quotas=scenario_quotas,
        scenario_counts=scenario_counts,
        dataset_fingerprint=fingerprint,
        generation_provenance=provenance,
        files=files,
        generation_time_seconds=round(time.perf_counter() - started, 3),
        generation_peak_memory_mb=resource_monitor.peak_memory_mb,
        generation_resource_estimate=resource_estimate,
    )
    write_json(dataset_path / "manifest.json", manifest)
    return manifest


def _write_docx(
    request: DatasetCreateRequest, docx_path: Path, pages: int
) -> tuple[
    list[FactRecord], list[QuestionRecord], list[DocumentObject], list[ObjectRelation]
]:
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError as exc:
        raise RuntimeError("python-docx is required to generate DOCX datasets") from exc

    rng = random.Random(request.seed)
    document = Document()
    document.core_properties.title = request.title
    document.core_properties.subject = "面向 LightRAG 的中文文档记忆测评"
    document.core_properties.keywords = "LightRAG, 中文, 测评, 检索, 标准答案"
    document.add_heading(request.title, 0)
    document.add_paragraph(
        "本文档为 LightRAG 中文文档记忆测评自动生成。文中包含可验证的事实、"
        "干扰信息、表格、图示、公式与跨文档证据，所有题目的标准答案均保存在数据集 oracle 中。"
    )

    facts: list[FactRecord] = []
    questions: list[QuestionRecord] = []
    objects: list[DocumentObject] = []
    relations: list[ObjectRelation] = []
    object_counter = 0
    relation_counter = 0

    def add_object(
        object_type: str,
        *,
        title: str,
        text: str,
        section: str,
        page: int,
        parent_id: str = "",
        labels: list[str] | None = None,
    ) -> str:
        nonlocal object_counter
        object_counter += 1
        object_id = f"OBJ-{object_counter:05d}"
        if request.profile == "rich":
            objects.append(
                DocumentObject(
                    object_id=object_id,
                    object_type=object_type,  # type: ignore[arg-type]
                    title=title,
                    text=text,
                    section=section,
                    page_start=page,
                    parent_id=parent_id,
                    labels=labels or [],
                )
            )
        return object_id

    def add_relation(
        source_id: str, target_id: str, relation_type: str, evidence: str = ""
    ) -> None:
        nonlocal relation_counter
        if request.profile != "rich":
            return
        relation_counter += 1
        relations.append(
            ObjectRelation(
                relation_id=f"REL-{relation_counter:05d}",
                source_id=source_id,
                target_id=target_id,
                relation_type=relation_type,  # type: ignore[arg-type]
                evidence_text=evidence,
            )
        )

    root_id = add_object(
        "document",
        title=request.title,
        text="中文合成测评文档及其证据对象图。",
        section="文档",
        page=1,
        labels=["document", "zh"],
    )
    fact_counter = 1
    previous_text_fact = ""
    previous_table_fact = ""
    last_delivery_fact = ""
    last_risk_fact = ""

    if request.profile == "rich":
        document.add_heading("项目背景与治理边界", level=1)
        context_text = (
            "“星桥客户服务知识升级”项目在 2026 年第三季度推进。项目负责人林岚负责业务验收，"
            "数据治理负责人周衡负责来源准入，安全负责人顾澄负责涉及隐私与退款场景的上线签核。"
            "任何批次只有同时满足质量门槛、责任人签核和回滚预案完整三个条件，才能进入生产环境。"
        )
        document.add_paragraph(context_text)
        context_id = add_object(
            "paragraph",
            title="项目背景",
            text=context_text,
            section="项目背景与治理边界",
            page=1,
            parent_id=root_id,
            labels=["program_context", "governance"],
        )
        add_relation(root_id, context_id, "contains", context_text)
        context_fact_id = f"FACT-{fact_counter:05d}"
        fact_counter += 1
        facts.append(
            FactRecord(
                fact_id=context_fact_id,
                fact_type="governance_owner",
                answer="林岚、周衡和顾澄分别负责业务验收、来源准入与安全签核",
                expected_text=context_text,
                section="项目背景与治理边界",
                page=1,
                object_type="text",
                object_id_hint=context_id,
            )
        )
        questions.append(
            QuestionRecord(
                id=f"Q-{context_fact_id}",
                question="星桥项目中，业务验收、来源准入和安全签核分别由谁负责？",
                answer="林岚、周衡和顾澄分别负责业务验收、来源准入与安全签核",
                question_type="multi_hop",
                evidence_fact_ids=[context_fact_id],
            )
        )

    figure_counter = 0
    for page in range(1, pages + 1):
        chapter = (page - 1) // 10 + 1
        unit = (page - 1) // 2 + 1
        workstream_title, _ = _WORKSTREAMS[
            (page - 1) % len(_WORKSTREAMS)
        ]
        chapter_title = f"第 {chapter} 章：星桥项目——{workstream_title}"
        section_title = f"第 {chapter}.{unit} 节：实施单元 {unit:04d}"
        section = f"{chapter_title} / {section_title} / 第 {page} 页"
        if page == 1 or (page - 1) % 10 == 0:
            document.add_heading(chapter_title, level=1)
        if page == 1 or (page - 1) % 2 == 0:
            document.add_heading(section_title, level=2)
        section_id = add_object(
            "section",
            title=section_title,
            text=f"本节记录{workstream_title}的目标、证据与决策约束。",
            section=section,
            page=page,
            parent_id=root_id,
            labels=["section"],
        )
        add_relation(root_id, section_id, "contains")

        param, param_unit, value_fn = _PARAMETER_PROFILES[
            (page - 1) % len(_PARAMETER_PROFILES)
        ]
        value = value_fn(page, rng)
        fact_id = f"FACT-{fact_counter:05d}"
        fact_counter += 1
        text = _SPEC_TEMPLATES[(page - 1) % len(_SPEC_TEMPLATES)].format(
            fact_id=fact_id,
            page=page,
            workstream=workstream_title,
            param=param,
            value=value,
            unit=param_unit,
        )
        document.add_heading(f"第 {page} 页规格说明", level=3)
        document.add_paragraph(text)
        paragraph_id = add_object(
            "paragraph",
            title=fact_id,
            text=text,
            section=section,
            page=page,
            parent_id=section_id,
            labels=["gold_fact"],
        )
        add_relation(section_id, paragraph_id, "contains", text)
        answer = f"{value} {param_unit}".strip()
        facts.append(
            FactRecord(
                fact_id=fact_id,
                fact_type="direct_numeric",
                answer=answer,
                expected_text=text,
                section=section,
                page=page,
                object_type="text",
                object_id_hint=paragraph_id if request.profile == "rich" else "",
            )
        )
        questions.append(
            QuestionRecord(
                id=f"Q-{fact_id}",
                question=f"第 {page} 页{workstream_title}的{param}是多少？",
                answer=answer,
                question_type="direct_numeric",
                evidence_fact_ids=[fact_id],
            )
        )
        previous_text_fact = fact_id

        narrative = _NARRATIVE_TEMPLATES[(page - 1) % len(_NARRATIVE_TEMPLATES)].format(
            workstream=workstream_title,
            param=param,
            value=value,
            unit=param_unit,
            page=page,
            prev_page=page - 1 if page > 1 else page,
        )
        document.add_paragraph(narrative)
        document.add_paragraph(
            _CONTROL_TEMPLATES[(page - 1) % len(_CONTROL_TEMPLATES)].format(
                param=param, value=value, unit=param_unit, workstream=workstream_title
            )
        )

        if page % 3 == 2:
            near = _near_value(value, param_unit, rng)
            document.add_paragraph(
                _DISTRACTOR_TEMPLATES[(page // 3) % len(_DISTRACTOR_TEMPLATES)].format(
                    near=near, unit=param_unit
                )
            )

        if request.profile == "rich":
            if page % 4 == 1:
                delivery_fact_id = f"FACT-{fact_counter:05d}"
                fact_counter += 1
                milestone = f"M-{chapter}{unit:02d}"
                delivery_text = (
                    f"{delivery_fact_id}：实施单元 {unit:04d} 的交付里程碑为 {milestone}，"
                    f"由林岚在第 {page + 2} 页对应的业务验收窗口确认。未完成验收的资料仅可停留在灰度环境。"
                )
                document.add_paragraph(delivery_text)
                delivery_id = add_object(
                    "paragraph",
                    title=delivery_fact_id,
                    text=delivery_text,
                    section=section,
                    page=page,
                    parent_id=section_id,
                    labels=["milestone", "gold_fact"],
                )
                add_relation(section_id, delivery_id, "contains", delivery_text)
                facts.append(
                    FactRecord(
                        fact_id=delivery_fact_id,
                        fact_type="delivery_milestone",
                        answer=milestone,
                        expected_text=delivery_text,
                        section=section,
                        page=page,
                        object_type="text",
                        object_id_hint=delivery_id,
                    )
                )
                last_delivery_fact = delivery_fact_id
            elif page % 4 == 3:
                risk_fact_id = f"FACT-{fact_counter:05d}"
                fact_counter += 1
                control = "周衡完成来源复核且顾澄完成安全签核"
                risk_variants = (
                    (
                        f"{risk_fact_id}：实施单元 {unit:04d} 涉及退款或隐私变更时，"
                        f"即使业务里程碑已达成，发布前仍必须由{control}。"
                    ),
                    (
                        f"{risk_fact_id}：若实施单元 {unit:04d} 触及退款或隐私信息，"
                        f"达到业务里程碑后还需满足“{control}”的双重控制条件，方可进入生产。"
                    ),
                    (
                        f"{risk_fact_id}：实施单元 {unit:04d} 的退款或隐私相关变更在里程碑达成后，"
                        f"仍须通过“{control}”的发布门禁。"
                    ),
                )
                risk_text = risk_variants[(page // 4) % len(risk_variants)]
                document.add_paragraph(risk_text)
                risk_id = add_object(
                    "paragraph",
                    title=risk_fact_id,
                    text=risk_text,
                    section=section,
                    page=page,
                    parent_id=section_id,
                    labels=["risk_control", "gold_fact"],
                )
                add_relation(section_id, risk_id, "contains", risk_text)
                facts.append(
                    FactRecord(
                        fact_id=risk_fact_id,
                        fact_type="release_constraint",
                        answer=control,
                        expected_text=risk_text,
                        section=section,
                        page=page,
                        object_type="text",
                        object_id_hint=risk_id,
                    )
                )
                last_risk_fact = risk_fact_id
            elif page % 4 == 0 and last_delivery_fact and last_risk_fact:
                questions.append(
                    QuestionRecord(
                        id=f"Q-RELEASE-GATE-{page:04d}",
                        question=f"依据第 {page} 页附近的交付里程碑和风险控制要求，"
                        "涉及退款或隐私信息的单元在满足里程碑后还需要什么才能上线？",
                        answer="周衡完成来源复核且顾澄完成安全签核",
                        question_type="multi_hop",
                        evidence_fact_ids=[last_delivery_fact, last_risk_fact],
                    )
                )

        if "tables" in request.modalities and page % 3 == 0:
            table_fact_id = f"FACT-{fact_counter:05d}"
            fact_counter += 1
            theme, question_template, table_unit, table_value_fn = _TABLE_THEMES[
                (page // 3) % len(_TABLE_THEMES)
            ]
            table_answer = table_value_fn(page, rng)
            caption = f"表 {page}：{workstream_title}的{theme}"
            document.add_paragraph(caption)
            if theme == "时延阈值":
                rows = [
                    ("参数", "标称值", "最大值 (ms)"),
                    ("延迟 Alpha", f"{page}.0", f"{page * 3}.5"),
                    ("延迟 Beta", f"{page + 2}.0", f"{page * 4}.5"),
                    ("重排序", f"{page * 0.5:.1f}", f"{page + 1}.0"),
                    ("生成首 token", f"{page * 0.8:.1f}", f"{page + 2}.0"),
                    (table_fact_id, "标准行标记", f"{table_answer} {table_unit}"),
                ]
            elif theme == "吞吐配额":
                rows = [
                    ("参数", "标称值", "配额 (次/秒)"),
                    ("检索并发", f"{page * 10}", f"{page * 16}"),
                    ("重排并发", f"{page * 6}", f"{page * 10}"),
                    ("批处理", f"{page * 4}", f"{page * 8}"),
                    ("缓存命中", f"{page * 30}", f"{page * 44}"),
                    (table_fact_id, "标准行标记", f"{table_answer} {table_unit}"),
                ]
            else:
                rows = [
                    ("参数", "标称值", "预算上限 (%)"),
                    ("检索错误", f"{page * 0.1:.1f}", f"{page * 0.2:.1f}"),
                    ("重排错误", f"{page * 0.05:.2f}", f"{page * 0.1:.2f}"),
                    ("生成错误", f"{page * 0.08:.2f}", f"{page * 0.16:.2f}"),
                    ("超时率", f"{page * 0.04:.2f}", f"{page * 0.08:.2f}"),
                    (table_fact_id, "标准行标记", f"{table_answer} {table_unit}"),
                ]
            table = document.add_table(rows=len(rows), cols=3)
            table.style = "Table Grid"
            for row_index, row_values in enumerate(rows):
                for column_index, value_text in enumerate(row_values):
                    table.rows[row_index].cells[column_index].text = str(value_text)
            table_id = add_object(
                "table",
                title=caption,
                text=f"{table_fact_id} 标准行标记 {table_answer} {table_unit}",
                section=section,
                page=page,
                parent_id=section_id,
                labels=["table", "gold_fact"],
            )
            add_relation(section_id, table_id, "contains")
            table_fact_answer = f"{table_answer} {table_unit}".strip()
            facts.append(
                FactRecord(
                    fact_id=table_fact_id,
                    fact_type="table_cell",
                    answer=table_fact_answer,
                    expected_text=f"{table_fact_id} 标准行标记 {table_fact_answer}",
                    section=section,
                    page=page,
                    object_type="table",
                    object_id_hint=table_id if request.profile == "rich" else caption,
                )
            )
            questions.append(
                QuestionRecord(
                    id=f"Q-{table_fact_id}",
                    question=question_template.format(page=page),
                    answer=table_fact_answer,
                    question_type="table_cell",
                    evidence_fact_ids=[table_fact_id],
                )
            )
            previous_table_fact = table_fact_id

        if "figures" in request.modalities and page % 4 == 0:
            figure_fact_id = f"FACT-{fact_counter:05d}"
            fact_counter += 1
            figure_counter += 1
            state = _FIGURE_STATES[(page // 4) % len(_FIGURE_STATES)]
            fig_path = docx_path.parent / f"zh_figure_{page:04d}.png"
            _write_figure(
                fig_path,
                kind=(page // 4) % 4,
                page=page,
                workstream=workstream_title,
                state=state,
                value=value,
                unit=param_unit,
            )
            document.add_picture(str(fig_path), width=Inches(4.6))
            caption = f"图 {page}：{figure_fact_id} 表示第 {page} 页{workstream_title}的{state}。"
            document.add_paragraph(caption)
            figure_id = add_object(
                "figure",
                title=f"图 {page}",
                text=caption,
                section=section,
                page=page,
                parent_id=section_id,
                labels=["figure", "caption"],
            )
            add_relation(section_id, figure_id, "contains", caption)
            figure_answer = f"第 {page} 页{workstream_title}的{state}"
            facts.append(
                FactRecord(
                    fact_id=figure_fact_id,
                    fact_type="figure_caption",
                    answer=figure_answer,
                    expected_text=caption,
                    section=section,
                    page=page,
                    object_type="figure",
                    object_id_hint=figure_id if request.profile == "rich" else f"图 {page}",
                )
            )
            questions.append(
                QuestionRecord(
                    id=f"Q-{figure_fact_id}",
                    question=f"图 {page} 表示的是什么状态？",
                    answer=figure_answer,
                    question_type="figure_caption",
                    evidence_fact_ids=[figure_fact_id],
                )
            )

        if "equations" in request.modalities and page % 5 == 0:
            equation_fact_id = f"FACT-{fact_counter:05d}"
            fact_counter += 1
            metric, formula, explanation = _EQUATION_POOL[
                (page // 5) % len(_EQUATION_POOL)
            ]
            equation_text = (
                f"{equation_fact_id}：第 {page} 页{workstream_title}的{metric}由公式 "
                f"{formula} 定义，其中{explanation}。"
            )
            document.add_paragraph(equation_text)
            equation_id = add_object(
                "equation",
                title=f"公式 {page}",
                text=equation_text,
                section=section,
                page=page,
                parent_id=section_id,
                labels=["equation", "gold_fact"],
            )
            add_relation(section_id, equation_id, "contains", equation_text)
            facts.append(
                FactRecord(
                    fact_id=equation_fact_id,
                    fact_type="equation",
                    answer=formula,
                    expected_text=equation_text,
                    section=section,
                    page=page,
                    object_type="equation",
                    object_id_hint=equation_id if request.profile == "rich" else f"公式 {page}",
                )
            )
            questions.append(
                QuestionRecord(
                    id=f"Q-{equation_fact_id}",
                    question=f"第 {page} 页{workstream_title}的{metric}由什么公式定义？",
                    answer=formula,
                    question_type="equation",
                    evidence_fact_ids=[equation_fact_id],
                )
            )
            if previous_table_fact and previous_text_fact:
                questions.append(
                    QuestionRecord(
                        id=f"Q-MULTIHOP-{page:04d}",
                        question=f"结合第 {page} 页公式和最近的标准表格行，应引用哪两个事实编号？",
                        answer=f"{previous_table_fact}；{equation_fact_id}",
                        question_type="multi_hop",
                        evidence_fact_ids=[previous_table_fact, equation_fact_id],
                    )
                )

        if page < pages:
            document.add_page_break()

    questions.append(
        QuestionRecord(
            id="Q-ABSTAIN-00001",
            question="文档中不存在的锆石旁路模块的审批编号是什么？",
            answer="文档未提供此信息。",
            question_type="abstain",
            evidence_fact_ids=[],
            expected_behavior="abstain",
        )
    )
    document.save(docx_path)
    return facts, questions, objects, relations


def _near_value(value: str, unit: str, rng: random.Random) -> str:
    """Return a plausible near-miss value for distractor sentences."""
    try:
        numeric = float(value)
    except (TypeError, ValueError):
        return value
    delta = max(abs(numeric) * 0.02, 1.0)
    near = numeric + delta * (1 if rng.random() < 0.5 else -1)
    if unit in {"%", "QMU"}:
        return f"{near:.2f}" if unit == "%" else f"{near:.0f}"
    return f"{near:.1f}"


def _write_artifacts(
    *,
    dataset_path: Path,
    dataset_id: str,
    facts: list[FactRecord],
    questions: list[QuestionRecord],
    objects: list[DocumentObject],
    relations: list[ObjectRelation],
    oracle: OraclePayload,
) -> None:
    common = {"dataset_id": dataset_id, "language": "zh"}
    write_json(dataset_path / "facts.json", {**common, "facts": facts})
    write_json(dataset_path / "questions.json", {**common, "questions": questions})
    write_json(dataset_path / "objects.json", {**common, "objects": objects})
    write_json(dataset_path / "relations.json", {**common, "relations": relations})
    write_json(dataset_path / "oracle.json", oracle)


def _stable_dataset_id(request: DatasetCreateRequest, pages: int) -> str:
    raw = (
        f"zh:{request.profile}:{request.tier}:{pages}:{request.formats}:"
        f"{request.modalities}:{request.seed}:{request.split}:"
        f"{sorted(request.scenario_quotas.items())}"
    )
    digest = hashlib.md5(raw.encode("utf-8")).hexdigest()[:10]
    return f"zh-{request.profile}-{request.tier}-{pages}p-{digest}"


def _load_cjk_font(size: int):
    from PIL import ImageFont

    candidates = (
        "/System/Library/Fonts/STHeiti Medium.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
        "/System/Library/Fonts/Hiragino Sans GB.ttc",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
    )
    for path in candidates:
        try:
            return ImageFont.truetype(path, size)
        except Exception as exc:  # noqa: BLE001 - try the next font candidate
            logger.debug("CJK font %s unavailable: %s", path, exc)
            continue
    return ImageFont.load_default()


def _write_figure(
    path: Path,
    *,
    kind: int,
    page: int,
    workstream: str,
    state: str,
    value: str,
    unit: str,
) -> None:
    """Draw a labelled diagram/chart for one figure page."""
    try:
        from PIL import Image, ImageDraw
    except ImportError as exc:
        raise RuntimeError("Pillow is required to generate figure assets") from exc

    width, height = 900, 420
    image = Image.new("RGB", (width, height), color=(245, 247, 250))
    draw = ImageDraw.Draw(image)
    title_font = _load_cjk_font(26)
    label_font = _load_cjk_font(20)
    small_font = _load_cjk_font(16)
    blue = (27, 99, 157)
    green = (35, 130, 91)
    orange = (198, 122, 33)
    ink = (60, 66, 74)

    draw.text((28, 18), f"第 {page} 页 · {workstream}", fill=ink, font=title_font)
    draw.text((28, 52), state, fill=blue, font=small_font)
    draw.text((28, 76), f"基准参数：{value} {unit}", fill=green, font=small_font)

    if kind == 0:
        # Flow diagram: 输入 -> 解析/分块 -> 检索/重排 -> 生成/引用
        boxes = ("请求输入", "解析与分块", "检索与重排", "生成与引用")
        xs = [40, 250, 460, 670]
        for index, label in enumerate(boxes):
            x = xs[index]
            fill = (232, 240, 248) if index != 2 else (233, 244, 238)
            outline = blue if index != 2 else green
            draw.rounded_rectangle((x, 200, x + 170, 300), radius=16, fill=fill, outline=outline, width=4)
            draw.text((x + 85 - 30, 240), label, fill=ink, font=label_font)
            if index < 3:
                x1 = x + 170
                y1 = 250
                x2 = xs[index + 1] - 10
                draw.line((x1, y1, x2, y1), fill=ink, width=4)
                draw.polygon([(x2 - 14, y1 - 8), (x2 + 4, y1), (x2 - 14, y1 + 8)], fill=ink)
    elif kind == 1:
        # Bar chart
        values = [max(4, int(page * 1.2)), max(6, int(page * 1.8)), max(5, int(page * 1.5)), max(7, int(page * 2.4))]
        labels = ("检索", "重排", "生成", "引用")
        x0, y0, bw = 90, 320, 150
        draw.line((x0, y0, x0, y0 - 230), fill=ink, width=3)
        draw.line((x0, y0, 850, y0), fill=ink, width=3)
        max_v = max(values) * 1.2
        for index, (v, label) in enumerate(zip(values, labels)):
            x = x0 + 30 + index * 190
            h = int(220 * v / max_v)
            color = (blue, green, orange, ink)[index]
            draw.rectangle((x, y0 - h, x + bw, y0), fill=color)
            draw.text((x + 50, y0 - h - 26), str(v), fill=ink, font=small_font)
            draw.text((x + 40, y0 + 12), label, fill=ink, font=label_font)
    elif kind == 2:
        # Line chart
        points = [(i, 40 + page * 6 + i * 22 + (i * i) % 13) for i in range(8)]
        x0, y0 = 90, 320
        draw.line((x0, y0, x0, y0 - 230), fill=ink, width=3)
        draw.line((x0, y0, 850, y0), fill=ink, width=3)
        max_y = max(p[1] for p in points) * 1.2
        coords = [(x0 + 90 + i * 90, y0 - int(210 * p[1] / max_y)) for i, p in enumerate(points)]
        draw.line(coords, fill=blue, width=4, joint="curve")
        for index, (x, y) in enumerate(coords):
            draw.ellipse((x - 6, y - 6, x + 6, y + 6), fill=orange, outline=ink, width=2)
            draw.text((x - 16, y + 8), str(points[index][1]), fill=ink, font=small_font)
        draw.text((60, 336), "时间窗口", fill=ink, font=label_font)
    else:
        # Layered blocks
        layers = (
            ("数据源与接入", blue),
            ("知识处理与索引", green),
            ("评测与质量门禁", orange),
            ("线上服务与审计", ink),
        )
        y = 180
        for label, color in layers:
            draw.rounded_rectangle((120, y, 760, y + 52), radius=12, fill=(240, 244, 248), outline=color, width=3)
            draw.text((150, y + 12), label, fill=ink, font=label_font)
            y += 64
    image.save(path)
