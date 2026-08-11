from __future__ import annotations

import hashlib
import random
import shutil
import subprocess
import time
from pathlib import Path

from memory_data_service.cross_document import add_cross_document_case
from memory_data_service.provenance import (
    annotate_question_scenarios,
    build_provenance,
    resolve_scenario_quotas,
)
from memory_data_service.resource_guard import (
    GenerationResourceMonitor,
    enforce_generation_limits,
    estimate_generation_resources,
)
from memory_data_service.schemas import (
    DatasetCreateRequest,
    DatasetManifest,
    FactRecord,
    GeneratedFile,
    OraclePayload,
    QuestionRecord,
)
from memory_data_service.storage import DEFAULT_GENERATED_ROOT, ensure_root, write_json

SOFFICE_DEFAULT = Path("/opt/homebrew/bin/soffice")


def generate_dataset(
    request: DatasetCreateRequest,
    *,
    root: Path = DEFAULT_GENERATED_ROOT,
) -> DatasetManifest:
    started = time.perf_counter()
    root = ensure_root(root)
    pages = request.resolved_pages()
    enforce_generation_limits(request, pages)
    resource_estimate = estimate_generation_resources(request, pages)
    dataset_id = request.dataset_id or _stable_dataset_id(request, pages)
    dataset_path = root / dataset_id
    dataset_path.mkdir(parents=True, exist_ok=True)

    docx_path = dataset_path / f"{dataset_id}.docx"
    with GenerationResourceMonitor() as resource_monitor:
        facts, questions = _write_docx(request, docx_path, pages)
        companion_docx = (
            add_cross_document_case(
                dataset_id=dataset_id,
                dataset_path=dataset_path,
                facts=facts,
                questions=questions,
                language=request.language,
            )
            if "docx" in request.formats
            else None
        )
        scenario_counts = annotate_question_scenarios(questions)
        scenario_quotas = resolve_scenario_quotas(
            requested=request.scenario_quotas, observed=scenario_counts
        )
        if "pdf" in request.formats:
            pdf_record = _convert_pdf(docx_path, dataset_path)
        else:
            pdf_record = _skipped("pdf")

        oracle = OraclePayload(
            dataset_id=dataset_id, language=request.language, facts=facts, questions=questions
        )
        write_json(
            dataset_path / "facts.json",
            {"dataset_id": dataset_id, "language": request.language, "facts": facts},
        )
        write_json(
            dataset_path / "questions.json",
            {"dataset_id": dataset_id, "language": request.language, "questions": questions},
        )
        write_json(dataset_path / "objects.json", {"dataset_id": dataset_id, "language": request.language, "objects": []})
        write_json(dataset_path / "relations.json", {"dataset_id": dataset_id, "language": request.language, "relations": []})
        write_json(dataset_path / "oracle.json", oracle)
    files: list[GeneratedFile] = [
        _file_record(docx_path, "docx", role="source_document")
        if "docx" in request.formats
        else _skipped("docx")
    ]
    if "pdf" in request.formats:
        files.append(pdf_record)
    else:
        files.append(_skipped("pdf"))
    if companion_docx is not None:
        files.append(_file_record(companion_docx, "docx", role="source_document"))
    for name in ("facts.json", "questions.json", "objects.json", "relations.json", "oracle.json"):
        files.append(_file_record(dataset_path / name, "json"))
    for asset_path in sorted(dataset_path.glob("*.png")):
        files.append(_file_record(asset_path, "png"))

    fingerprint, provenance = build_provenance(
        request=request,
        pages=pages,
        generator="basic_docx_generator",
        template_version="basic-docx-v1",
        source_file=Path(__file__),
    )
    manifest = DatasetManifest(
        dataset_id=dataset_id,
        tier=request.tier,
        pages=pages,
        profile=request.profile,
        language=request.language,
        formats=request.formats,
        modalities=request.modalities,
        title=request.title,
        display_name=request.display_name,
        split=request.split,
        scenario_quotas=scenario_quotas,
        scenario_counts=scenario_counts,
        dataset_fingerprint=fingerprint,
        generation_provenance=provenance,
        files=files,
        generation_time_seconds=round(time.perf_counter() - started, 3),
        generation_peak_memory_mb=resource_monitor.peak_memory_mb,
        generation_resource_estimate=resource_estimate,
    )
    write_json(dataset_path / "manifest.json", manifest)
    return manifest


def _stable_dataset_id(request: DatasetCreateRequest, pages: int) -> str:
    raw = (
        f"{request.tier}:{pages}:{request.formats}:{request.modalities}:{request.seed}:"
        f"{request.split}:{sorted(request.scenario_quotas.items())}"
    )
    digest = hashlib.md5(raw.encode("utf-8")).hexdigest()[:10]
    return f"{request.tier}-{pages}p-{digest}"


def _write_docx(
    request: DatasetCreateRequest,
    docx_path: Path,
    pages: int,
) -> tuple[list[FactRecord], list[QuestionRecord]]:
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError as exc:
        raise RuntimeError("python-docx is required to generate DOCX datasets") from exc

    rng = random.Random(request.seed)
    doc = Document()
    doc.core_properties.title = request.title
    doc.add_heading(request.title, 0)
    doc.add_paragraph(
        "This synthetic document is generated for LightRAG single-document memory "
        "evaluation. It intentionally contains facts, distractors, tables, figures, "
        "equations, captions, and cross references."
    )

    image_path = docx_path.parent / "synthetic_figure.png"
    if "figures" in request.modalities:
        _write_placeholder_image(image_path)

    facts: list[FactRecord] = []
    questions: list[QuestionRecord] = []
    fact_counter = 1

    for page in range(1, pages + 1):
        h1 = f"System Area {((page - 1) // 25) + 1}"
        h2 = f"Subsystem {((page - 1) // 5) + 1}"
        section = f"{h1} / {h2} / Page {page}"
        if page == 1 or (page - 1) % 25 == 0:
            doc.add_heading(h1, level=1)
        if page == 1 or (page - 1) % 5 == 0:
            doc.add_heading(h2, level=2)
        doc.add_heading(f"Page {page} Specification", level=3)

        direct_fact_id = f"FACT-{fact_counter:05d}"
        value = 1000 + page * 7 + rng.randint(0, 6)
        fact_counter += 1
        direct_sentence = (
            f"{direct_fact_id}: The canonical calibration limit for page {page} "
            f"is {value} QMU. Distractor limit values {value + 11} QMU and "
            f"{value - 9} QMU are obsolete and must not be used."
        )
        doc.add_paragraph(direct_sentence)
        facts.append(
            FactRecord(
                fact_id=direct_fact_id,
                fact_type="direct_numeric",
                answer=f"{value} QMU",
                expected_text=direct_sentence,
                section=section,
                page=page,
                object_type="text",
            )
        )
        questions.append(
            QuestionRecord(
                id=f"Q-{direct_fact_id}",
                question=f"What is the canonical calibration limit for page {page}?",
                answer=f"{value} QMU",
                question_type="direct_numeric",
                evidence_fact_ids=[direct_fact_id],
            )
        )

        for paragraph_index in range(3):
            doc.add_paragraph(
                "The subsystem narrative repeats nearby terms to stress retrieval. "
                f"Page {page} paragraph {paragraph_index + 1} discusses control loops, "
                "memory alignment, auditability, source tracing, and reference hygiene."
            )

        if "tables" in request.modalities and page % 5 == 0:
            table_fact_id = f"FACT-{fact_counter:05d}"
            fact_counter += 1
            table_answer = f"{page * 3}.5 ms"
            doc.add_paragraph(f"Table {page}: Timing thresholds for subsystem {h2}.")
            table = doc.add_table(rows=4, cols=3)
            table.style = "Table Grid"
            rows = [
                ("Parameter", "Nominal", "Maximum"),
                ("Latency Alpha", f"{page}.0 ms", table_answer),
                ("Latency Beta", f"{page + 2}.0 ms", f"{page * 4}.5 ms"),
                (table_fact_id, "Gold row marker", table_answer),
            ]
            for row_idx, row_values in enumerate(rows):
                for col_idx, value_text in enumerate(row_values):
                    table.rows[row_idx].cells[col_idx].text = value_text
            facts.append(
                FactRecord(
                    fact_id=table_fact_id,
                    fact_type="table_cell",
                    answer=table_answer,
                    expected_text=f"{table_fact_id} Gold row marker {table_answer}",
                    section=section,
                    page=page,
                    object_type="table",
                    object_id_hint=f"Table {page}",
                )
            )
            questions.append(
                QuestionRecord(
                    id=f"Q-{table_fact_id}",
                    question=f"In Table {page}, what is the Maximum value for the gold row marker?",
                    answer=table_answer,
                    question_type="table_cell",
                    evidence_fact_ids=[table_fact_id],
                )
            )

        if "figures" in request.modalities and page % 6 == 0:
            figure_fact_id = f"FACT-{fact_counter:05d}"
            fact_counter += 1
            doc.add_picture(str(image_path), width=Inches(3.8))
            caption = (
                f"Figure {page}: {figure_fact_id} shows the verified control-flow "
                f"state for page {page}."
            )
            doc.add_paragraph(caption)
            facts.append(
                FactRecord(
                    fact_id=figure_fact_id,
                    fact_type="figure_caption",
                    answer=f"verified control-flow state for page {page}",
                    expected_text=caption,
                    section=section,
                    page=page,
                    object_type="figure",
                    object_id_hint=f"Figure {page}",
                )
            )

        if "equations" in request.modalities and page % 4 == 0:
            equation_fact_id = f"FACT-{fact_counter:05d}"
            fact_counter += 1
            equation = f"{equation_fact_id}: E_{{{page}}} = P_{{{page}}} * T_{{{page}}} / eta"
            doc.add_paragraph(equation)
            facts.append(
                FactRecord(
                    fact_id=equation_fact_id,
                    fact_type="equation",
                    answer=f"E_{page} = P_{page} * T_{page} / eta",
                    expected_text=equation,
                    section=section,
                    page=page,
                    object_type="equation",
                    object_id_hint=f"Equation {page}",
                )
            )
            questions.append(
                QuestionRecord(
                    id=f"Q-{equation_fact_id}",
                    question=f"What equation defines E_{page}?",
                    answer=f"E_{page} = P_{page} * T_{page} / eta",
                    question_type="equation",
                    evidence_fact_ids=[equation_fact_id],
                )
            )

        if page < pages:
            doc.add_page_break()

    questions.append(
        QuestionRecord(
            id="Q-ABSTAIN-00001",
            question="What is the approval code for the nonexistent zirconium bypass module?",
            answer="The document does not provide this information.",
            question_type="abstain",
            evidence_fact_ids=[],
            expected_behavior="abstain",
        )
    )

    doc.save(docx_path)
    return facts, questions


def _write_placeholder_image(path: Path) -> None:
    try:
        from PIL import Image, ImageDraw
    except ImportError as exc:
        raise RuntimeError("Pillow is required to generate figure assets") from exc

    image = Image.new("RGB", (900, 420), color=(245, 247, 250))
    draw = ImageDraw.Draw(image)
    draw.rectangle((60, 80, 300, 220), outline=(20, 96, 160), width=5)
    draw.rectangle((560, 80, 800, 220), outline=(20, 140, 96), width=5)
    draw.line((300, 150, 560, 150), fill=(80, 80, 80), width=4)
    draw.polygon([(545, 135), (575, 150), (545, 165)], fill=(80, 80, 80))
    draw.text((95, 130), "Input State", fill=(20, 96, 160))
    draw.text((590, 130), "Verified State", fill=(20, 140, 96))
    image.save(path)


def _convert_pdf(docx_path: Path, dataset_path: Path) -> GeneratedFile:
    soffice = SOFFICE_DEFAULT if SOFFICE_DEFAULT.exists() else shutil.which("soffice")
    if not soffice:
        return _skipped("pdf", "LibreOffice/soffice not found")
    try:
        subprocess.run(
            [
                str(soffice),
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(dataset_path),
                str(docx_path),
            ],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            timeout=180,
        )
    except Exception as exc:  # pragma: no cover - depends on local LibreOffice
        return _skipped("pdf", f"PDF conversion failed: {exc}")
    pdf_path = docx_path.with_suffix(".pdf")
    if not pdf_path.exists():
        return _skipped("pdf", "PDF conversion did not produce an output file")
    return _file_record(pdf_path, "pdf", role="source_document")


def _file_record(
    path: Path, fmt: str, *, role: str = "evaluation_artifact"
) -> GeneratedFile:
    return GeneratedFile(
        name=path.name,
        format=fmt,  # type: ignore[arg-type]
        role=role,  # type: ignore[arg-type]
        path=str(path),
        size_bytes=path.stat().st_size if path.exists() else 0,
    )


def _skipped(fmt: str, message: str = "not requested") -> GeneratedFile:
    return GeneratedFile(
        name=f"skipped.{fmt}",
        format=fmt,  # type: ignore[arg-type]
        role="source_document" if fmt in {"docx", "pdf"} else "evaluation_artifact",
        path="",
        size_bytes=0,
        status="skipped",
        message=message,
    )
