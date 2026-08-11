from __future__ import annotations

import hashlib
import random
import re
import shutil
import time
import zipfile
from pathlib import Path
from xml.sax.saxutils import escape

from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches

from memory_data_service.cross_document import add_cross_document_case
from memory_data_service.generators.docx_generator import (
    _convert_pdf,
    _file_record,
    _skipped,
)
from memory_data_service.provenance import (
    annotate_question_scenarios,
    build_provenance,
    resolve_scenario_quotas,
)
from memory_data_service.resource_guard import (
    GenerationResourceMonitor,
    enforce_generation_limits,
    estimate_generation_resources,
)
from memory_data_service.schemas import (
    DatasetCreateRequest,
    DatasetManifest,
    DocumentObject,
    FactRecord,
    GeneratedFile,
    ObjectRelation,
    OraclePayload,
    QuestionRecord,
)
from memory_data_service.storage import DEFAULT_GENERATED_ROOT, ensure_root, write_json

MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"


class RichDocxBuilder:
    def __init__(self, request: DatasetCreateRequest, dataset_id: str, dataset_path: Path):
        self.request = request
        self.dataset_id = dataset_id
        self.dataset_path = dataset_path
        self.rng = random.Random(request.seed)
        self.objects: list[DocumentObject] = []
        self.relations: list[ObjectRelation] = []
        self.facts: list[FactRecord] = []
        self.questions: list[QuestionRecord] = []
        self.footnotes: list[tuple[int, str]] = []
        self.endnotes: list[tuple[int, str]] = []
        self.fact_counter = 1
        self.object_counter = 1
        self.relation_counter = 1
        self.bookmark_counter = 1
        self.last_delivery_fact = ""
        self.last_risk_fact = ""

    def build(self, docx_path: Path, pages: int) -> None:
        from docx import Document

        doc = Document()
        doc.core_properties.title = self.request.title
        doc.core_properties.subject = "Synthetic benchmark for single-document memory evaluation"
        doc.core_properties.keywords = "LightRAG, memory, rich document, oracle"
        self._configure_document(doc)

        root_id = self._add_object(
            "document",
            title=self.request.title,
            text="Synthetic rich technical document with oracle object graph.",
            section="document",
            page_start=1,
            page_end=pages,
        )

        self._cover(doc)
        self._toc(doc, pages)
        self._executive_summary(doc, root_id)
        self._add_program_dossier(doc, root_id)
        self._add_note_references(doc, root_id)
        self._add_complex_layout_controls(doc, root_id)

        last_table_fact = ""
        last_text_fact = ""
        last_figure_fact = ""
        section_ids: dict[str, str] = {}

        for page in range(1, pages + 1):
            chapter = ((page - 1) // 10) + 1
            section_index = ((page - 1) // 2) + 1
            chapter_title = f"Chapter {chapter}: Adaptive Retrieval Control Area"
            section_title = f"Section {chapter}.{section_index}: Retrieval Cell Group {section_index:04d}"
            section_path = f"{chapter_title} / {section_title}"

            if page == 1 or (page - 1) % 10 == 0:
                doc.add_page_break()
                doc.add_heading(chapter_title, level=1)
                chapter_id = self._add_object(
                    "section",
                    title=chapter_title,
                    section=chapter_title,
                    page_start=page,
                    parent_id=root_id,
                    labels=["chapter"],
                )
                section_ids[chapter_title] = chapter_id
                self._add_relation(root_id, chapter_id, "contains")

            if page == 1 or (page - 1) % 2 == 0:
                doc.add_heading(section_title, level=2)
                parent_id = section_ids[chapter_title]
                section_id = self._add_object(
                    "section",
                    title=section_title,
                    section=section_path,
                    page_start=page,
                    parent_id=parent_id,
                    labels=["section"],
                )
                section_ids[section_title] = section_id
                self._add_relation(parent_id, section_id, "contains")
                self._add_section_summary(doc, page, section_path, section_id)
            else:
                section_id = section_ids[section_title]

            text_fact_id, paragraph_id = self._add_gold_paragraph(
                doc, page, section_path, section_id
            )
            last_text_fact = text_fact_id

            self._add_distractor_paragraph(doc, page, section_path, section_id, text_fact_id)
            self._add_numbered_controls(doc, page, section_path, section_id)
            self._add_operational_dependency(doc, page, section_path, section_id)

            if page % 7 == 0:
                self._add_version_fact(doc, page, section_path, section_id)

            if page % 9 == 0:
                self._add_conflict_fact(doc, page, section_path, section_id)

            if page % 11 == 0:
                self._add_negative_constraint_fact(doc, page, section_path, section_id)

            if "tables" in self.request.modalities and page % 3 == 0:
                last_table_fact = self._add_rich_table(doc, page, section_path, section_id)

            if "figures" in self.request.modalities and page % 4 == 0:
                last_figure_fact = self._add_rich_figure(doc, page, section_path, section_id)

            if "equations" in self.request.modalities and page % 5 == 0:
                equation_fact_id = self._add_omml_equation(doc, page, section_path, section_id)
                if last_table_fact:
                    self._add_multihop_question(
                        page=page,
                        question_id=f"Q-MULTIHOP-{page:04d}",
                        fact_ids=[last_table_fact, equation_fact_id],
                        question=(
                            f"Using the latest timing table before page {page} and Equation EQ-{page:04d}, "
                            f"which latency fact and equation should be cited together?"
                        ),
                        answer=f"{self._fact_answer(last_table_fact)}; {self._fact_answer(equation_fact_id)}",
                    )

            if page % 6 == 0 and last_table_fact and last_text_fact:
                self._add_table_reference_paragraph(doc, page, section_path, section_id, last_table_fact)
                self._add_multihop_question(
                    page=page,
                    question_id=f"Q-CROSS-{page:04d}",
                    fact_ids=[last_text_fact, last_table_fact],
                    question=(
                        f"For retrieval cell {page:04d}, combine the calibration limit "
                        "with the nearest preceding table maximum."
                    ),
                    answer=f"{self._fact_answer(last_text_fact)}; {self._fact_answer(last_table_fact)}",
                )

            if page % 8 == 0 and last_figure_fact:
                self._add_reference_paragraph(doc, page, section_path, section_id, last_figure_fact)
                if last_text_fact:
                    self._add_multihop_question(
                        page=page,
                        question_id=f"Q-FIGTEXT-{page:04d}",
                        fact_ids=[last_text_fact, last_figure_fact],
                        question=(
                            f"For Retrieval Cell {page:04d}, combine the authoritative "
                            "calibration limit with the cited figure control state."
                        ),
                        answer=f"{self._fact_answer(last_text_fact)}; {self._fact_answer(last_figure_fact)}",
                        question_type="figure_text",
                    )

            self._add_local_conclusion(doc, page, section_path, section_id)

            if page < pages:
                doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        self._glossary(doc, root_id, pages)
        self._references(doc, root_id, pages)
        self._appendix(doc, root_id, pages)
        self._abstention_questions()
        doc.save(docx_path)
        self._inject_note_parts(docx_path)
        self._inject_update_fields_setting(docx_path)

    def _configure_document(self, doc) -> None:
        section = doc.sections[0]
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.header.paragraphs[0].text = "LightRAG Memory Benchmark - Synthetic Rich Document"
        section.footer.paragraphs[0].text = "Confidential synthetic evaluation corpus"

    def _cover(self, doc) -> None:
        doc.add_heading(self.request.title, 0)
        doc.add_paragraph("Document class: controlled operational decision dossier.")
        doc.add_paragraph(f"Dataset id: {self.dataset_id}")
        doc.add_paragraph(
            "The Northstar Knowledge Operations programme is preparing a phased customer-support "
            "release. This dossier records business outcomes, accountable owners, quality gates, "
            "risk controls, and the technical evidence required for each decision."
        )

    def _toc(self, doc, pages: int) -> None:
        doc.add_page_break()
        doc.add_heading("Table of Contents", level=1)
        self._add_toc_field(doc)
        for chapter in range(1, max(2, (pages + 9) // 10) + 1):
            doc.add_paragraph(f"Chapter {chapter}: Adaptive Retrieval Control Area")
            for section_offset in range(1, 6):
                section_no = (chapter - 1) * 5 + section_offset
                doc.add_paragraph(
                    f"Section {chapter}.{section_no}: Retrieval Cell {section_no:04d}",
                    style="List Bullet",
                )

    def _add_toc_field(self, doc) -> None:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instruction = OxmlElement("w:instrText")
        instruction.set(qn("xml:space"), "preserve")
        instruction.text = 'TOC \\o "1-3" \\h \\z \\u'
        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        placeholder = OxmlElement("w:t")
        placeholder.text = "Right-click and update field to render the automatic TOC."
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.extend([begin, instruction, separate, placeholder, end])

    def _executive_summary(self, doc, root_id: str) -> None:
        doc.add_page_break()
        doc.add_heading("Executive Summary", level=1)
        paragraph = (
            "Northstar must reduce repeat customer contacts without allowing unreviewed policy "
            "or privacy-sensitive guidance into production. Each release therefore connects a "
            "business owner, a data steward, a security approver, measurable quality gates, "
            "and a rollback decision. The benchmark evaluates whether a retrieval system can "
            "preserve these dependencies across narrative text, tables, visual captions, "
            "formulas, references, and deliberately plausible drafts."
        )
        doc.add_paragraph(paragraph)
        paragraph_id = self._add_object(
            "paragraph",
            text=paragraph,
            section="Executive Summary",
            page_start=1,
            parent_id=root_id,
            labels=["summary"],
        )
        self._add_relation(root_id, paragraph_id, "contains")

    def _add_program_dossier(self, doc, root_id: str) -> None:
        doc.add_heading("Programme Governance and Decision Rights", level=1)
        text = (
            "FACT-GOV-00001: In the Northstar programme, Maya Chen owns business acceptance, "
            "Darius Holt owns source-data approval, and Priya Shah owns security sign-off. "
            "A release may enter production only when the quality gate, source approval, "
            "security sign-off, and rollback plan are all complete."
        )
        doc.add_paragraph(text)
        object_id = self._add_object(
            "paragraph",
            title="FACT-GOV-00001",
            text=text,
            section="Programme Governance and Decision Rights",
            page_start=1,
            parent_id=root_id,
            labels=["programme_context", "governance", "gold_fact"],
        )
        self._add_relation(root_id, object_id, "contains")
        self._add_fact(
            fact_id="FACT-GOV-00001",
            fact_type="governance_owner",
            answer="Maya Chen: business acceptance; Darius Holt: source-data approval; Priya Shah: security sign-off",
            expected_text=text,
            section="Programme Governance and Decision Rights",
            page=1,
            object_type="text",
            object_id_hint=object_id,
        )
        self._add_question(
            id="Q-FACT-GOV-00001",
            question="Who owns business acceptance, source-data approval, and security sign-off in Northstar?",
            answer="Maya Chen: business acceptance; Darius Holt: source-data approval; Priya Shah: security sign-off",
            question_type="multi_hop",
            evidence_fact_ids=["FACT-GOV-00001"],
        )
        self._add_relation(object_id, "FACT-GOV-00001", "supports", evidence_text=text)

    def _add_operational_dependency(self, doc, page: int, section: str, section_id: str) -> None:
        """Create interdependent delivery and risk evidence across nearby pages."""
        if page % 4 == 1:
            fact_id = self._next_fact_id()
            milestone = f"NS-GATE-{page:02d}"
            text = (
                f"{fact_id}: Delivery commitment for Retrieval Cell {page:04d}: "
                f"{milestone} is the business-acceptance gate owned by Maya Chen. "
                "Materials that miss this gate may remain in controlled preview but cannot be released."
            )
            doc.add_paragraph(text)
            object_id = self._add_object(
                "paragraph",
                title=fact_id,
                text=text,
                section=section,
                page_start=page,
                parent_id=section_id,
                labels=["delivery_milestone", "gold_fact"],
            )
            self._add_relation(section_id, object_id, "contains")
            self._add_fact(
                fact_id=fact_id,
                fact_type="delivery_milestone",
                answer=milestone,
                expected_text=text,
                section=section,
                page=page,
                object_type="text",
                object_id_hint=object_id,
            )
            self.last_delivery_fact = fact_id
        elif page % 4 == 3:
            fact_id = self._next_fact_id()
            control = "Darius Holt source approval and Priya Shah security sign-off"
            text = (
                f"{fact_id}: Release constraint for Retrieval Cell {page:04d}: whenever the "
                f"change touches refunds or personal data, {control} are required after the "
                "business gate and before production deployment."
            )
            doc.add_paragraph(text)
            object_id = self._add_object(
                "paragraph",
                title=fact_id,
                text=text,
                section=section,
                page_start=page,
                parent_id=section_id,
                labels=["release_constraint", "gold_fact"],
            )
            self._add_relation(section_id, object_id, "contains")
            self._add_fact(
                fact_id=fact_id,
                fact_type="release_constraint",
                answer=control,
                expected_text=text,
                section=section,
                page=page,
                object_type="text",
                object_id_hint=object_id,
            )
            self.last_risk_fact = fact_id
        elif page % 4 == 0 and self.last_delivery_fact and self.last_risk_fact:
            self._add_multihop_question(
                page=page,
                question_id=f"Q-RELEASE-GATE-{page:04d}",
                fact_ids=[self.last_delivery_fact, self.last_risk_fact],
                question=(
                    "After the most recent business-acceptance gate is met, what additional "
                    "approvals are required before a refund or personal-data change can enter production?"
                ),
                answer="Darius Holt source approval and Priya Shah security sign-off",
            )

    def _add_note_references(self, doc, root_id: str) -> None:
        footnote_text = (
            "Footnote FN-0001: Source tracing must preserve note bodies separately "
            "from the paragraph that carries the note marker."
        )
        endnote_text = (
            "Endnote EN-0001: Endnote evidence is deliberately stored at document end "
            "to test nonlocal note retrieval."
        )
        paragraph = doc.add_paragraph("Document note controls include ")
        paragraph.add_run("a footnote marker")
        self._append_note_reference(paragraph, note_id=1, note_kind="footnote")
        paragraph.add_run(" and ")
        paragraph.add_run("an endnote marker")
        self._append_note_reference(paragraph, note_id=1, note_kind="endnote")
        paragraph.add_run(".")
        note_anchor_id = self._add_object(
            "paragraph",
            text="Document note controls include a footnote marker and an endnote marker.",
            section="Executive Summary",
            page_start=1,
            parent_id=root_id,
            labels=["note_anchor"],
        )
        footnote_id = self._add_object(
            "footnote",
            title="FN-0001",
            text=footnote_text,
            section="Executive Summary",
            page_start=1,
            parent_id=note_anchor_id,
            labels=["footnote"],
            metadata={"note_id": 1},
        )
        endnote_id = self._add_object(
            "endnote",
            title="EN-0001",
            text=endnote_text,
            section="Executive Summary",
            page_start=1,
            parent_id=note_anchor_id,
            labels=["endnote"],
            metadata={"note_id": 1},
        )
        self.footnotes.append((1, footnote_text))
        self.endnotes.append((1, endnote_text))
        self._add_relation(root_id, note_anchor_id, "contains")
        self._add_relation(note_anchor_id, footnote_id, "refers_to", evidence_text=footnote_text)
        self._add_relation(note_anchor_id, endnote_id, "refers_to", evidence_text=endnote_text)

    def _add_complex_layout_controls(self, doc, root_id: str) -> None:
        doc.add_page_break()
        section = doc.add_section(WD_SECTION_START.CONTINUOUS)
        self._set_column_count(section, 2)
        doc.add_heading("Complex Layout Controls", level=1)
        layout_id = self._add_object(
            "layout_region",
            title="Complex Layout Controls",
            text="Two-column synthetic section with a floating text box.",
            section="Complex Layout Controls",
            page_start=1,
            parent_id=root_id,
            labels=["complex_layout", "two_column_layout"],
            metadata={"columns": 2},
        )
        self._add_relation(root_id, layout_id, "contains")
        left_text = (
            "Column control COL-0001: the left column states that source-aware retrieval "
            "must preserve captions and object ids together."
        )
        right_text = (
            "Column control COL-0002: the right column states that layout-sensitive "
            "answers must not infer page coordinates unless parser evidence provides them."
        )
        for text, label in ((left_text, "left_column"), (right_text, "right_column")):
            doc.add_paragraph(text)
            paragraph_id = self._add_object(
                "paragraph",
                text=text,
                section="Complex Layout Controls",
                page_start=1,
                parent_id=layout_id,
                labels=["column_text", label],
            )
            self._add_relation(layout_id, paragraph_id, "contains")
        textbox_text = (
            "TEXTBOX-0001: Floating callout evidence says anchored objects must remain "
            "linked to their surrounding section."
        )
        self._add_textbox_paragraph(doc, textbox_text)
        textbox_id = self._add_object(
            "textbox",
            title="TEXTBOX-0001",
            text=textbox_text,
            section="Complex Layout Controls",
            page_start=1,
            parent_id=layout_id,
            labels=["textbox", "floating_object", "anchored_callout"],
            metadata={"shape": "vml_textbox"},
        )
        self._add_relation(layout_id, textbox_id, "contains")
        section = doc.add_section(WD_SECTION_START.CONTINUOUS)
        self._set_column_count(section, 1)

    def _set_column_count(self, section, count: int) -> None:
        sect_pr = section._sectPr
        cols = next((child for child in sect_pr if child.tag == qn("w:cols")), None)
        if cols is None:
            cols = OxmlElement("w:cols")
            sect_pr.append(cols)
        cols.set(qn("w:num"), str(count))

    def _add_textbox_paragraph(self, doc, text: str):
        paragraph = doc.add_paragraph()
        textbox_xml = f"""
        <w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
             xmlns:v="urn:schemas-microsoft-com:vml">
          <w:pict>
            <v:shape id="TextboxSynthetic01" type="#_x0000_t202"
              style="position:absolute;margin-left:0;margin-top:0;width:360pt;height:52pt;z-index:251659264">
              <v:textbox inset="5pt,5pt,5pt,5pt">
                <w:txbxContent>
                  <w:p>
                    <w:r><w:t>{escape(text)}</w:t></w:r>
                  </w:p>
                </w:txbxContent>
              </v:textbox>
            </v:shape>
          </w:pict>
        </w:r>
        """
        paragraph._p.append(parse_xml(textbox_xml))
        return paragraph

    def _append_note_reference(self, paragraph, *, note_id: int, note_kind: str) -> None:
        run = paragraph.add_run()
        element_name = "w:footnoteReference" if note_kind == "footnote" else "w:endnoteReference"
        reference = OxmlElement(element_name)
        reference.set(qn("w:id"), str(note_id))
        run._r.append(reference)

    def _add_section_summary(self, doc, page: int, section: str, section_id: str) -> None:
        text = (
            f"Section summary {page:04d}: this section binds authoritative retrieval-cell "
            "facts to nearby table, figure, formula, and reference objects."
        )
        doc.add_paragraph(text)
        summary_id = self._add_object(
            "paragraph",
            text=text,
            section=section,
            page_start=page,
            parent_id=section_id,
            labels=["section_summary"],
        )
        self._add_relation(section_id, summary_id, "contains")

    def _add_local_conclusion(self, doc, page: int, section: str, section_id: str) -> None:
        text = (
            f"Local conclusion {page:04d}: current answers must prefer gold FACT rows, "
            "caption-grounded visual states, and explicit equation labels over archived distractors."
        )
        doc.add_paragraph(text)
        conclusion_id = self._add_object(
            "paragraph",
            text=text,
            section=section,
            page_start=page,
            parent_id=section_id,
            labels=["local_conclusion"],
        )
        self._add_relation(section_id, conclusion_id, "contains")

    def _add_gold_paragraph(
        self,
        doc,
        page: int,
        section: str,
        section_id: str,
    ) -> tuple[str, str]:
        fact_id = self._next_fact_id()
        limit = 9000 + page * 17 + self.rng.randint(0, 9)
        unit = "QMU"
        text = (
            f"{fact_id}: The authoritative calibration limit for Retrieval Cell {page:04d} "
            f"is {limit} {unit}. This value supersedes the provisional value "
            f"{limit + 13} {unit} and the legacy value {limit - 21} {unit}."
        )
        doc.add_paragraph(text)
        paragraph_id = self._add_object(
            "paragraph",
            text=text,
            section=section,
            page_start=page,
            parent_id=section_id,
            labels=["gold_fact", "numeric"],
        )
        self._add_relation(section_id, paragraph_id, "contains")
        self._add_fact(
            fact_id=fact_id,
            fact_type="direct_numeric",
            answer=f"{limit} {unit}",
            expected_text=text,
            section=section,
            page=page,
            object_type="text",
            object_id_hint=paragraph_id,
        )
        self._add_question(
            id=f"Q-{fact_id}",
            question=f"What is the authoritative calibration limit for Retrieval Cell {page:04d}?",
            answer=f"{limit} {unit}",
            question_type="direct_numeric",
            evidence_fact_ids=[fact_id],
        )
        self._add_relation(paragraph_id, fact_id, "supports", evidence_text=text)
        return fact_id, paragraph_id

    def _add_distractor_paragraph(
        self,
        doc,
        page: int,
        section: str,
        section_id: str,
        target_fact_id: str,
    ) -> None:
        text = (
            f"DISTRACTOR-{page:04d}: Historical audit logs mention Retrieval Cell {page:04d} "
            f"near terms such as calibration, limit, threshold, and QMU, but those values "
            "belong to an archived scenario and must not be used for current answers."
        )
        doc.add_paragraph(text)
        paragraph_id = self._add_object(
            "paragraph",
            text=text,
            section=section,
            page_start=page,
            parent_id=section_id,
            labels=["distractor"],
        )
        self._add_relation(section_id, paragraph_id, "contains")
        self._add_relation(paragraph_id, target_fact_id, "distracts", evidence_text=text)

    def _add_numbered_controls(self, doc, page: int, section: str, section_id: str) -> None:
        intro = f"Operational checklist for Retrieval Cell {page:04d}:"
        doc.add_paragraph(intro)
        paragraph_id = self._add_object(
            "paragraph",
            text=intro,
            section=section,
            page_start=page,
            parent_id=section_id,
            labels=["checklist"],
        )
        self._add_relation(section_id, paragraph_id, "contains")
        for item in (
            "Use the authoritative calibration limit, not provisional values.",
            "Prefer table gold rows over archived rows.",
            "When figure and caption disagree, caption evidence takes precedence.",
        ):
            doc.add_paragraph(item, style="List Number")
            item_id = self._add_object(
                "paragraph",
                text=item,
                section=section,
                page_start=page,
                parent_id=paragraph_id,
                labels=["list_item", "numbered_control"],
            )
            self._add_relation(paragraph_id, item_id, "contains")
        bullet_items = (
            (
                "Evidence bundle",
                "Combine section text with table, figure, and equation objects before final answers.",
            ),
            (
                "Traceability bundle",
                "Preserve block ids, captions, and reference notes for citation audits.",
            ),
        )
        for parent_text, child_text in bullet_items:
            self._add_styled_paragraph(doc, parent_text, "List Bullet")
            parent_item_id = self._add_object(
                "paragraph",
                text=parent_text,
                section=section,
                page_start=page,
                parent_id=paragraph_id,
                labels=["list_item", "bullet_control"],
            )
            self._add_relation(paragraph_id, parent_item_id, "contains")
            self._add_styled_paragraph(doc, child_text, "List Bullet 2")
            child_item_id = self._add_object(
                "paragraph",
                text=child_text,
                section=section,
                page_start=page,
                parent_id=parent_item_id,
                labels=["list_item", "nested_bullet_control"],
            )
            self._add_relation(parent_item_id, child_item_id, "contains")

    def _add_styled_paragraph(self, doc, text: str, style: str) -> None:
        try:
            doc.add_paragraph(text, style=style)
        except KeyError:
            doc.add_paragraph(text)

    def _add_version_fact(self, doc, page: int, section: str, section_id: str) -> str:
        fact_id = self._next_fact_id()
        version = f"v{page}.{self.rng.randint(1, 4)}"
        effective = f"2026-Q{((page - 1) % 4) + 1}"
        text = (
            f"{fact_id}: For Retrieval Cell {page:04d}, policy version {version} is active "
            f"from {effective}. Earlier policy version v{page - 1}.9 is expired and must be ignored."
        )
        doc.add_paragraph(text)
        object_id = self._add_object(
            "paragraph",
            text=text,
            section=section,
            page_start=page,
            parent_id=section_id,
            labels=["version_condition", "gold_fact"],
        )
        self._add_relation(section_id, object_id, "contains")
        self._add_fact(
            fact_id=fact_id,
            fact_type="version_condition",
            answer=f"{version} from {effective}",
            expected_text=text,
            section=section,
            page=page,
            object_type="text",
            object_id_hint=object_id,
        )
        self._add_question(
            id=f"Q-{fact_id}",
            question=f"Which policy version is active for Retrieval Cell {page:04d}, and from when?",
            answer=f"{version} from {effective}",
            question_type="version_condition",
            evidence_fact_ids=[fact_id],
        )
        self._add_relation(object_id, fact_id, "supports", evidence_text=text)
        return fact_id

    def _add_conflict_fact(self, doc, page: int, section: str, section_id: str) -> str:
        fact_id = self._next_fact_id()
        canonical_method = f"Method-C{page:04d}"
        obsolete_method = f"Method-D{page:04d}"
        text = (
            f"{fact_id}: Conflict resolution for Retrieval Cell {page:04d}: "
            f"{canonical_method} is canonical because it is signed by the audit board; "
            f"{obsolete_method} is a conflicting draft and must not be cited as final."
        )
        doc.add_paragraph(text)
        object_id = self._add_object(
            "paragraph",
            text=text,
            section=section,
            page_start=page,
            parent_id=section_id,
            labels=["conflict_resolution", "gold_fact"],
        )
        self._add_relation(section_id, object_id, "contains")
        draft_text = (
            f"CONFLICT-DRAFT-{page:04d}: {obsolete_method} appears in a draft memo, "
            "but it conflicts with the audit-board-signed canonical method."
        )
        doc.add_paragraph(draft_text)
        draft_obj_id = self._add_object(
            "paragraph",
            text=draft_text,
            section=section,
            page_start=page,
            parent_id=section_id,
            labels=["conflicting_draft", "distractor"],
        )
        self._add_relation(section_id, draft_obj_id, "contains")
        self._add_fact(
            fact_id=fact_id,
            fact_type="conflict_resolution",
            answer=canonical_method,
            expected_text=text,
            section=section,
            page=page,
            object_type="text",
            object_id_hint=object_id,
        )
        self._add_question(
            id=f"Q-{fact_id}",
            question=f"When methods conflict for Retrieval Cell {page:04d}, which method is canonical?",
            answer=canonical_method,
            question_type="conflict_resolution",
            evidence_fact_ids=[fact_id],
        )
        self._add_relation(object_id, fact_id, "supports", evidence_text=text)
        self._add_relation(draft_obj_id, fact_id, "contradicts", evidence_text=draft_text)
        return fact_id

    def _add_negative_constraint_fact(self, doc, page: int, section: str, section_id: str) -> str:
        fact_id = self._next_fact_id()
        forbidden = f"retired override channel ROC-{page:04d}"
        text = (
            f"{fact_id}: Negative constraint for Retrieval Cell {page:04d}: "
            f"the {forbidden} must not be used in any current answer, even when it appears "
            "near calibration or latency terminology."
        )
        doc.add_paragraph(text)
        object_id = self._add_object(
            "paragraph",
            text=text,
            section=section,
            page_start=page,
            parent_id=section_id,
            labels=["negative_constraint", "gold_fact"],
        )
        self._add_relation(section_id, object_id, "contains")
        self._add_fact(
            fact_id=fact_id,
            fact_type="negative_constraint",
            answer=forbidden,
            expected_text=text,
            section=section,
            page=page,
            object_type="text",
            object_id_hint=object_id,
        )
        self._add_question(
            id=f"Q-{fact_id}",
            question=f"Which override channel must not be used for Retrieval Cell {page:04d}?",
            answer=forbidden,
            question_type="negative_constraint",
            evidence_fact_ids=[fact_id],
        )
        self._add_relation(object_id, fact_id, "supports", evidence_text=text)
        return fact_id

    def _add_rich_table(self, doc, page: int, section: str, section_id: str) -> str:
        fact_id = self._next_fact_id()
        table_id = f"TBL-{page:04d}"
        caption = f"Table {table_id}: Consolidated latency thresholds for Retrieval Cell {page:04d}."
        self._add_bookmarked_paragraph(doc, caption, _bookmark_name(table_id))
        caption_id = self._add_object(
            "caption",
            title=caption,
            text=caption,
            section=section,
            page_start=page,
            parent_id=section_id,
            labels=["table_caption"],
        )
        table = doc.add_table(rows=7, cols=5)
        table.style = "Table Grid"
        table.cell(0, 0).text = "Subsystem"
        table.cell(0, 1).merge(table.cell(0, 2)).text = "Nominal Band"
        table.cell(0, 3).merge(table.cell(0, 4)).text = "Safety Band"
        headers = ["Parameter", "Minimum", "Nominal", "Maximum", "Unit"]
        for col, header in enumerate(headers):
            table.cell(1, col).text = header
        max_value = page * 11 + 0.75
        rows = [
            ("Latency Alpha", f"{page + 1}.00", f"{page + 2}.50", f"{max_value:.2f}", "ms"),
            ("Latency Beta", f"{page + 2}.00", f"{page + 3}.50", f"{max_value + 8:.2f}", "ms"),
            ("Voltage Guard", f"{page + 0.1:.1f}", f"{page + 0.5:.1f}", f"{page + 0.9:.1f}", "V"),
            (fact_id, "gold-row", "authoritative", f"{max_value:.2f}", "ms"),
            ("Archived Row", "legacy", "obsolete", f"{max_value + 17:.2f}", "ms"),
        ]
        for row_offset, row_values in enumerate(rows, start=2):
            for col, value in enumerate(row_values):
                table.cell(row_offset, col).text = str(value)
        footnote = (
            f"Table note {table_id}: The gold-row maximum is authoritative; the Archived Row is a distractor."
        )
        doc.add_paragraph(footnote)

        table_id_obj = self._add_object(
            "table",
            title=table_id,
            text=f"{caption} {fact_id} gold-row authoritative {max_value:.2f} ms",
            section=section,
            page_start=page,
            parent_id=section_id,
            labels=["table", "merged_header", "gold_fact"],
            metadata={"rows": 7, "columns": 5},
        )
        self._add_relation(section_id, caption_id, "contains")
        self._add_relation(section_id, table_id_obj, "contains")
        self._add_relation(caption_id, table_id_obj, "caption_of", evidence_text=caption)
        self._add_fact(
            fact_id=fact_id,
            fact_type="table_cell",
            answer=f"{max_value:.2f} ms",
            expected_text=f"{fact_id} gold-row authoritative {max_value:.2f} ms",
            section=section,
            page=page,
            object_type="table",
            object_id_hint=table_id_obj,
        )
        self._add_question(
            id=f"Q-{fact_id}",
            question=f"In {table_id}, what is the Maximum value for the authoritative gold row?",
            answer=f"{max_value:.2f} ms",
            question_type="table_cell",
            evidence_fact_ids=[fact_id],
        )
        self._add_relation(table_id_obj, fact_id, "supports")
        return fact_id

    def _add_rich_figure(self, doc, page: int, section: str, section_id: str) -> str:
        fact_id = self._next_fact_id()
        figure_id = f"FIG-{page:04d}"
        state = f"verified-state-{page:04d}"
        image_path = self.dataset_path / f"{figure_id}.png"
        self._write_information_figure(image_path, figure_id, state, fact_id, page)
        doc.add_picture(str(image_path), width=Inches(5.6))
        caption = (
            f"Figure {figure_id}: {fact_id} identifies the visual control state as {state}; "
            f"the nearby gray retired-state-{page:04d} label is obsolete."
        )
        doc.add_paragraph(caption)
        figure_obj_id = self._add_object(
            "figure",
            title=figure_id,
            text=f"{figure_id} {state} {fact_id}",
            section=section,
            page_start=page,
            parent_id=section_id,
            labels=["figure", "image_text", "gold_fact"],
            metadata={"asset": image_path.name},
        )
        caption_id = self._add_object(
            "caption",
            title=f"Figure caption {figure_id}",
            text=caption,
            section=section,
            page_start=page,
            parent_id=figure_obj_id,
            labels=["figure_caption"],
        )
        self._add_relation(section_id, figure_obj_id, "contains")
        self._add_relation(caption_id, figure_obj_id, "caption_of", evidence_text=caption)
        self._add_fact(
            fact_id=fact_id,
            fact_type="figure_caption",
            answer=state,
            expected_text=caption,
            section=section,
            page=page,
            object_type="figure",
            object_id_hint=figure_obj_id,
        )
        self._add_question(
            id=f"Q-{fact_id}",
            question=f"According to Figure {figure_id}, what is the visual control state?",
            answer=state,
            question_type="figure_caption",
            evidence_fact_ids=[fact_id],
        )
        self._add_relation(figure_obj_id, fact_id, "supports")
        return fact_id

    def _add_omml_equation(self, doc, page: int, section: str, section_id: str) -> str:
        fact_id = self._next_fact_id()
        variable_fact_id = self._next_fact_id()
        equation_id = f"EQ-{page:04d}"
        latex = f"E_{{{page}}}=P_{{{page}}}T_{{{page}}}/\\eta_{{{page}}}"
        doc.add_paragraph(f"Equation {equation_id}:")
        paragraph = doc.add_paragraph()
        paragraph._p.append(_omml_para(f"E_{page}=P_{page}T_{page}/eta_{page}"))
        definition = (
            f"Variable note {equation_id}: eta_{page} is the efficiency coefficient for "
            f"Retrieval Cell {page:04d}; P_{page} is the peak power and T_{page} is dwell time."
        )
        latex_mirror = f"LaTeX mirror {equation_id}: `{latex}`."
        doc.add_paragraph(definition)
        doc.add_paragraph(latex_mirror)
        doc.add_paragraph(
            f"Reference note {equation_id}: later calculations must cite Equation {equation_id} "
            "together with its variable note."
        )
        equation_obj_id = self._add_object(
            "equation",
            title=equation_id,
            text=latex,
            section=section,
            page_start=page,
            parent_id=section_id,
            labels=["omml", "equation", "gold_fact"],
        )
        self._add_relation(section_id, equation_obj_id, "contains")
        definition_obj_id = self._add_object(
            "paragraph",
            text=definition,
            section=section,
            page_start=page,
            parent_id=section_id,
            labels=["equation_variable_definition"],
        )
        self._add_relation(section_id, definition_obj_id, "contains")
        latex_obj_id = self._add_object(
            "paragraph",
            text=latex_mirror,
            section=section,
            page_start=page,
            parent_id=section_id,
            labels=["latex_formula_text", "formula_mirror"],
        )
        reference_obj_id = self._add_object(
            "reference",
            title=f"Reference note {equation_id}",
            text=(
                f"Reference note {equation_id}: later calculations must cite Equation {equation_id} "
                "together with its variable note."
            ),
            section=section,
            page_start=page,
            parent_id=section_id,
            labels=["equation_reference"],
        )
        self._add_relation(section_id, latex_obj_id, "contains")
        self._add_relation(section_id, reference_obj_id, "contains")
        self._add_relation(latex_obj_id, equation_obj_id, "mentions", evidence_text=latex_mirror)
        self._add_relation(reference_obj_id, equation_obj_id, "refers_to")
        self._add_fact(
            fact_id=fact_id,
            fact_type="equation",
            answer=latex,
            expected_text=f"Equation {equation_id}: {latex}",
            section=section,
            page=page,
            object_type="equation",
            object_id_hint=equation_obj_id,
        )
        self._add_question(
            id=f"Q-{fact_id}",
            question=f"What is the formula stated in Equation {equation_id}?",
            answer=latex,
            question_type="equation",
            evidence_fact_ids=[fact_id],
        )
        self._add_fact(
            fact_id=variable_fact_id,
            fact_type="equation_variable",
            answer=f"eta_{page} is the efficiency coefficient",
            expected_text=definition,
            section=section,
            page=page,
            object_type="equation",
            object_id_hint=definition_obj_id,
        )
        self._add_question(
            id=f"Q-{variable_fact_id}",
            question=f"In Equation {equation_id}, what does eta_{page} mean?",
            answer=f"eta_{page} is the efficiency coefficient",
            question_type="equation_variable",
            evidence_fact_ids=[variable_fact_id],
        )
        self._add_multihop_question(
            page=page,
            question_id=f"Q-EQVAR-{page:04d}",
            fact_ids=[fact_id, variable_fact_id],
            question=f"State Equation {equation_id} and define eta_{page}.",
            answer=f"{latex}; eta_{page} is the efficiency coefficient",
            question_type="formula_variable",
        )
        self._add_relation(equation_obj_id, fact_id, "supports")
        self._add_relation(definition_obj_id, variable_fact_id, "supports", evidence_text=definition)
        self._add_relation(definition_obj_id, equation_obj_id, "defines", evidence_text=definition)
        return fact_id

    def _add_table_reference_paragraph(
        self,
        doc,
        page: int,
        section: str,
        section_id: str,
        target_fact_id: str,
    ) -> None:
        ref_id = f"REF-TABLE-{page:04d}"
        text = (
            f"{ref_id}: The local acceptance review refers to table evidence {target_fact_id}; "
            "answers about latency maxima must cite that table row instead of narrative summaries."
        )
        paragraph = doc.add_paragraph(text + " Cross-reference field: ")
        self._append_ref_field(paragraph, _bookmark_name(f"TBL-{page - (page % 3):04d}"))
        ref_obj_id = self._add_object(
            "reference",
            title=ref_id,
            text=text,
            section=section,
            page_start=page,
            parent_id=section_id,
            labels=["table_cross_reference"],
        )
        self._add_relation(section_id, ref_obj_id, "contains")
        self._add_relation(ref_obj_id, target_fact_id, "refers_to", evidence_text=text)

    def _add_reference_paragraph(
        self,
        doc,
        page: int,
        section: str,
        section_id: str,
        target_fact_id: str,
    ) -> None:
        ref_id = f"REF-{page:04d}"
        text = (
            f"{ref_id}: The implementation review must cite the figure evidence for "
            f"{target_fact_id} before accepting a visual-state answer."
        )
        doc.add_paragraph(text)
        ref_obj_id = self._add_object(
            "reference",
            title=ref_id,
            text=text,
            section=section,
            page_start=page,
            parent_id=section_id,
            labels=["cross_reference"],
        )
        self._add_relation(section_id, ref_obj_id, "contains")
        self._add_relation(ref_obj_id, target_fact_id, "refers_to", evidence_text=text)

    def _glossary(self, doc, root_id: str, pages: int) -> None:
        doc.add_page_break()
        doc.add_heading("Glossary", level=1)
        glossary_id = self._add_object(
            "section",
            title="Glossary",
            section="Glossary",
            page_start=pages + 1,
            parent_id=root_id,
            labels=["glossary"],
        )
        self._add_relation(root_id, glossary_id, "contains")
        terms = {
            "Authoritative calibration limit": "The current value that supersedes all provisional and legacy limits.",
            "Object hit rate": "The proportion of retrieved evidence that preserves the target table, figure, or equation object.",
            "Abstention": "A refusal when no document evidence supports the requested answer.",
        }
        for term, definition in terms.items():
            text = f"{term}: {definition}"
            doc.add_paragraph(text)
            term_id = self._add_object(
                "glossary_term",
                title=term,
                text=definition,
                section="Glossary",
                page_start=pages + 1,
                parent_id=glossary_id,
                labels=["definition"],
            )
            self._add_relation(glossary_id, term_id, "contains")
            self._add_relation(term_id, glossary_id, "defines", evidence_text=text)

    def _references(self, doc, root_id: str, pages: int) -> None:
        doc.add_page_break()
        doc.add_heading("References", level=1)
        references_id = self._add_object(
            "section",
            title="References",
            section="References",
            page_start=pages + 2,
            parent_id=root_id,
            labels=["references"],
        )
        self._add_relation(root_id, references_id, "contains")
        entries = [
            "[REF-STD-001] Synthetic Retrieval Control Standard, revision 2026-A.",
            "[REF-TBL-014] Table preservation guidance for merged-header benchmark documents.",
            "[REF-FIG-021] Visual-state traceability protocol for caption-grounded QA.",
        ]
        citation_paragraph = doc.add_paragraph("Citation field control: ")
        self._append_simple_field(
            citation_paragraph,
            " CITATION REFSTD001 \\l 1033 ",
            "[REF-STD-001]",
        )
        citation_id = self._add_object(
            "reference",
            title="CITATION-REFSTD001",
            text="Citation field control: [REF-STD-001]",
            section="References",
            page_start=pages + 2,
            parent_id=references_id,
            labels=["citation_field", "bibliographic_reference"],
        )
        self._add_relation(references_id, citation_id, "contains")
        for entry in entries:
            doc.add_paragraph(entry)
            ref_id = self._add_object(
                "reference",
                title=entry.split("]", 1)[0].lstrip("["),
                text=entry,
                section="References",
                page_start=pages + 2,
                parent_id=references_id,
                labels=["bibliographic_reference"],
            )
            self._add_relation(references_id, ref_id, "contains")
        bibliography_paragraph = doc.add_paragraph("Bibliography field control: ")
        self._append_simple_field(
            bibliography_paragraph,
            " BIBLIOGRAPHY \\l 1033 ",
            "Synthetic bibliography field placeholder.",
        )
        bibliography_id = self._add_object(
            "reference",
            title="BIBLIOGRAPHY-FIELD",
            text="Bibliography field control: Synthetic bibliography field placeholder.",
            section="References",
            page_start=pages + 2,
            parent_id=references_id,
            labels=["bibliography_field", "bibliographic_reference"],
        )
        self._add_relation(references_id, bibliography_id, "contains")

    def _appendix(self, doc, root_id: str, pages: int) -> None:
        section = doc.add_section(WD_SECTION_START.NEW_PAGE)
        section.header.paragraphs[0].text = "Appendix - Synthetic distractor controls"
        doc.add_heading("Appendix A: Retired Scenario Catalog", level=1)
        appendix_id = self._add_object(
            "appendix",
            title="Appendix A: Retired Scenario Catalog",
            text="Retired scenario values used as controlled distractors.",
            section="Appendix A",
            page_start=pages + 2,
            parent_id=root_id,
            labels=["appendix", "distractor"],
        )
        self._add_relation(root_id, appendix_id, "contains")
        self._add_cross_page_table(doc, pages, appendix_id)
        for idx in range(1, 6):
            doc.add_paragraph(
                f"Appendix distractor A.{idx}: archived calibration limit {7000 + idx * 31} QMU "
                "is retained for negative retrieval testing only."
            )

    def _add_cross_page_table(self, doc, pages: int, appendix_id: str) -> None:
        fact_id = self._next_fact_id()
        table_id = "LONG-TBL-APP"
        answer = f"{pages + 42}.99 ms"
        caption = (
            f"Table {table_id}: Appendix long-table stress case with many rows; "
            f"{fact_id} marks the authoritative final rollover latency."
        )
        doc.add_paragraph(caption)
        caption_id = self._add_object(
            "caption",
            title=caption,
            text=caption,
            section="Appendix A",
            page_start=pages + 2,
            parent_id=appendix_id,
            labels=["table_caption", "long_table_caption"],
        )
        table = doc.add_table(rows=90, cols=4)
        table.style = "Table Grid"
        headers = ["Row", "Scenario", "Latency", "Status"]
        for col, header in enumerate(headers):
            table.cell(0, col).text = header
        for row_index in range(1, 90):
            table.cell(row_index, 0).text = f"A-{row_index:03d}"
            table.cell(row_index, 1).text = "appendix rollover stress"
            table.cell(row_index, 2).text = (
                answer if row_index == 89 else f"{pages + row_index / 10:.2f} ms"
            )
            table.cell(row_index, 3).text = fact_id if row_index == 89 else "distractor"
        table_obj_id = self._add_object(
            "table",
            title=table_id,
            text=f"{caption} {fact_id} authoritative final rollover latency {answer}",
            section="Appendix A",
            page_start=pages + 2,
            parent_id=appendix_id,
            labels=["table", "cross_page_long_table", "gold_fact"],
            metadata={"rows": 90, "columns": 4},
        )
        self._add_relation(appendix_id, caption_id, "contains")
        self._add_relation(appendix_id, table_obj_id, "contains")
        self._add_relation(caption_id, table_obj_id, "caption_of", evidence_text=caption)
        self._add_fact(
            fact_id=fact_id,
            fact_type="table_cell",
            answer=answer,
            expected_text=f"{fact_id} authoritative final rollover latency {answer}",
            section="Appendix A",
            page=pages + 2,
            object_type="table",
            object_id_hint=table_obj_id,
        )
        self._add_question(
            id=f"Q-{fact_id}",
            question=f"In {table_id}, what is the authoritative final rollover latency?",
            answer=answer,
            question_type="table_cell",
            evidence_fact_ids=[fact_id],
        )
        self._add_relation(table_obj_id, fact_id, "supports")

    def _abstention_questions(self) -> None:
        self._add_question(
            id="Q-ABSTAIN-00001",
            question="What is the approval code for the nonexistent zirconium bypass module?",
            answer="The document does not provide this information.",
            question_type="abstain",
            evidence_fact_ids=[],
            expected_behavior="abstain",
        )
        self._add_question(
            id="Q-ABSTAIN-00002",
            question="Which appendix authorizes the quantum coolant override?",
            answer="The document does not provide this information.",
            question_type="abstain",
            evidence_fact_ids=[],
            expected_behavior="abstain",
        )

    def _write_information_figure(
        self,
        path: Path,
        figure_id: str,
        state: str,
        fact_id: str,
        page: int,
    ) -> None:
        from PIL import Image, ImageDraw

        image = Image.new("RGB", (1200, 620), color=(250, 251, 253))
        draw = ImageDraw.Draw(image)
        draw.rectangle((40, 40, 1160, 580), outline=(50, 50, 60), width=3)
        draw.text((70, 65), f"{figure_id} control-state trace", fill=(20, 20, 30))
        draw.rectangle((90, 160, 390, 310), outline=(40, 110, 190), width=5)
        draw.text((125, 220), f"{fact_id}", fill=(40, 110, 190))
        draw.text((125, 250), state, fill=(40, 110, 190))
        draw.rectangle((780, 160, 1080, 310), outline=(150, 150, 150), width=5)
        draw.text((815, 220), f"retired-state-{page:04d}", fill=(150, 150, 150))
        draw.line((390, 235, 780, 235), fill=(80, 80, 90), width=4)
        draw.polygon([(760, 215), (800, 235), (760, 255)], fill=(80, 80, 90))
        draw.text((430, 390), f"Gold state: {state}", fill=(10, 120, 80))
        draw.text((430, 430), "Gray retired state is a distractor", fill=(130, 130, 130))
        image.save(path)

    def _add_multihop_question(
        self,
        *,
        page: int,
        question_id: str,
        fact_ids: list[str],
        question: str,
        answer: str,
        question_type: str = "multi_hop",
    ) -> None:
        self._add_question(
            id=question_id,
            question=question,
            answer=answer,
            question_type=question_type,
            evidence_fact_ids=fact_ids,
        )

    def _next_fact_id(self) -> str:
        fact_id = f"FACT-{self.fact_counter:05d}"
        self.fact_counter += 1
        return fact_id

    def _add_fact(self, **kwargs) -> None:
        self.facts.append(FactRecord(**kwargs))

    def _add_question(self, **kwargs) -> None:
        self.questions.append(QuestionRecord(**kwargs))

    def _add_object(self, object_type: str, **kwargs) -> str:
        object_id = kwargs.pop("object_id", f"OBJ-{self.object_counter:06d}")
        self.object_counter += 1
        self.objects.append(DocumentObject(object_id=object_id, object_type=object_type, **kwargs))
        return object_id

    def _add_relation(
        self,
        source_id: str,
        target_id: str,
        relation_type: str,
        *,
        evidence_text: str = "",
    ) -> None:
        relation_id = f"REL-{self.relation_counter:06d}"
        self.relation_counter += 1
        self.relations.append(
            ObjectRelation(
                relation_id=relation_id,
                source_id=source_id,
                target_id=target_id,
                relation_type=relation_type,
                evidence_text=evidence_text,
            )
        )

    def _fact_answer(self, fact_id: str) -> str:
        for fact in self.facts:
            if fact.fact_id == fact_id:
                return fact.answer
        return ""

    def _add_bookmarked_paragraph(self, doc, text: str, bookmark_name: str):
        paragraph = doc.add_paragraph()
        bookmark_id = str(self.bookmark_counter)
        self.bookmark_counter += 1
        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), bookmark_id)
        start.set(qn("w:name"), bookmark_name)
        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), bookmark_id)
        paragraph._p.append(start)
        paragraph.add_run(text)
        paragraph._p.append(end)
        return paragraph

    def _append_ref_field(self, paragraph, bookmark_name: str) -> None:
        self._append_simple_field(paragraph, f" REF {bookmark_name} \\h ", bookmark_name)

    def _append_simple_field(self, paragraph, instruction_text: str, placeholder_text: str) -> None:
        run = paragraph.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instruction = OxmlElement("w:instrText")
        instruction.set(qn("xml:space"), "preserve")
        instruction.text = instruction_text
        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        placeholder = OxmlElement("w:t")
        placeholder.text = placeholder_text
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.extend([begin, instruction, separate, placeholder, end])

    def _inject_update_fields_setting(self, docx_path: Path) -> None:
        tmp_path = docx_path.with_suffix(".settings.tmp.docx")
        with zipfile.ZipFile(docx_path, "r") as source, zipfile.ZipFile(
            tmp_path,
            "w",
            compression=zipfile.ZIP_DEFLATED,
        ) as target:
            for item in source.infolist():
                data = source.read(item.filename)
                if item.filename == "word/settings.xml":
                    settings_xml = data.decode("utf-8")
                    if "w:updateFields" not in settings_xml:
                        settings_xml = settings_xml.replace(
                            "</w:settings>",
                            '<w:updateFields w:val="true"/></w:settings>',
                        )
                    data = settings_xml.encode("utf-8")
                target.writestr(item, data)
        shutil.move(str(tmp_path), str(docx_path))

    def _inject_note_parts(self, docx_path: Path) -> None:
        if not self.footnotes and not self.endnotes:
            return
        tmp_path = docx_path.with_suffix(".notes.tmp.docx")
        with zipfile.ZipFile(docx_path, "r") as source, zipfile.ZipFile(
            tmp_path,
            "w",
            compression=zipfile.ZIP_DEFLATED,
        ) as target:
            content_types = source.read("[Content_Types].xml").decode("utf-8")
            rels = source.read("word/_rels/document.xml.rels").decode("utf-8")
            content_types = _ensure_override(
                content_types,
                "/word/footnotes.xml",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml",
            )
            content_types = _ensure_override(
                content_types,
                "/word/endnotes.xml",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.endnotes+xml",
            )
            rels = _ensure_relationship(
                rels,
                "rIdMemoryFootnotes",
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes",
                "footnotes.xml",
            )
            rels = _ensure_relationship(
                rels,
                "rIdMemoryEndnotes",
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes",
                "endnotes.xml",
            )
            skip = {
                "[Content_Types].xml",
                "word/_rels/document.xml.rels",
                "word/footnotes.xml",
                "word/endnotes.xml",
            }
            for item in source.infolist():
                if item.filename not in skip:
                    target.writestr(item, source.read(item.filename))
            target.writestr("[Content_Types].xml", content_types)
            target.writestr("word/_rels/document.xml.rels", rels)
            target.writestr("word/footnotes.xml", _notes_xml("footnotes", "footnote", self.footnotes))
            target.writestr("word/endnotes.xml", _notes_xml("endnotes", "endnote", self.endnotes))
        shutil.move(str(tmp_path), str(docx_path))


def generate_rich_dataset(
    request: DatasetCreateRequest,
    *,
    root: Path = DEFAULT_GENERATED_ROOT,
) -> DatasetManifest:
    started = time.perf_counter()
    root = ensure_root(root)
    pages = request.resolved_pages()
    enforce_generation_limits(request, pages)
    resource_estimate = estimate_generation_resources(request, pages)
    dataset_id = request.dataset_id or _stable_dataset_id(request, pages)
    dataset_path = root / dataset_id
    dataset_path.mkdir(parents=True, exist_ok=True)

    docx_path = dataset_path / f"{dataset_id}.docx"
    builder = RichDocxBuilder(request, dataset_id, dataset_path)
    with GenerationResourceMonitor() as resource_monitor:
        builder.build(docx_path, pages)
        companion_docx = (
            add_cross_document_case(
                dataset_id=dataset_id,
                dataset_path=dataset_path,
                facts=builder.facts,
                questions=builder.questions,
                language=request.language,
            )
            if "docx" in request.formats
            else None
        )
        scenario_counts = annotate_question_scenarios(builder.questions)
        scenario_quotas = resolve_scenario_quotas(
            requested=request.scenario_quotas, observed=scenario_counts
        )
        if "pdf" in request.formats:
            pdf_record = _convert_pdf(docx_path, dataset_path)
        else:
            pdf_record = _skipped("pdf")

        oracle = OraclePayload(
            dataset_id=dataset_id,
            language=request.language,
            facts=builder.facts,
            questions=builder.questions,
            objects=builder.objects,
            relations=builder.relations,
        )
        write_json(dataset_path / "facts.json", {"dataset_id": dataset_id, "language": request.language, "facts": builder.facts})
        write_json(
            dataset_path / "questions.json",
            {"dataset_id": dataset_id, "language": request.language, "questions": builder.questions},
        )
        write_json(
            dataset_path / "objects.json",
            {"dataset_id": dataset_id, "language": request.language, "objects": builder.objects},
        )
        write_json(
            dataset_path / "relations.json",
            {"dataset_id": dataset_id, "language": request.language, "relations": builder.relations},
        )
        write_json(dataset_path / "oracle.json", oracle)
    files: list[GeneratedFile] = [
        _file_record(docx_path, "docx", role="source_document")
        if "docx" in request.formats
        else _skipped("docx")
    ]
    files.append(pdf_record)
    if companion_docx is not None:
        files.append(_file_record(companion_docx, "docx", role="source_document"))
    for name in ("facts.json", "questions.json", "objects.json", "relations.json", "oracle.json"):
        files.append(_file_record(dataset_path / name, "json"))
    for asset_path in sorted(dataset_path.glob("*.png")):
        files.append(_file_record(asset_path, "png"))

    fingerprint, provenance = build_provenance(
        request=request,
        pages=pages,
        generator="rich_docx_generator",
        template_version="rich-docx-v2",
        source_file=Path(__file__),
    )
    manifest = DatasetManifest(
        dataset_id=dataset_id,
        tier=request.tier,
        pages=pages,
        profile=request.profile,
        language=request.language,
        formats=request.formats,
        modalities=request.modalities,
        title=request.title,
        display_name=request.display_name,
        split=request.split,
        scenario_quotas=scenario_quotas,
        scenario_counts=scenario_counts,
        dataset_fingerprint=fingerprint,
        generation_provenance=provenance,
        files=files,
        generation_time_seconds=round(time.perf_counter() - started, 3),
        generation_peak_memory_mb=resource_monitor.peak_memory_mb,
        generation_resource_estimate=resource_estimate,
    )
    write_json(dataset_path / "manifest.json", manifest)
    return manifest


def _stable_dataset_id(request: DatasetCreateRequest, pages: int) -> str:
    raw = (
        f"{request.profile}:{request.tier}:{pages}:{request.formats}:"
        f"{request.modalities}:{request.seed}:{request.split}:"
        f"{sorted(request.scenario_quotas.items())}"
    )
    digest = hashlib.md5(raw.encode("utf-8")).hexdigest()[:10]
    return f"{request.profile}-{request.tier}-{pages}p-{digest}"


def _bookmark_name(raw: str) -> str:
    return re.sub(r"[^A-Za-z0-9_]", "_", raw)


def _ensure_override(content_types: str, part_name: str, content_type: str) -> str:
    if f'PartName="{part_name}"' in content_types:
        return content_types
    override = f'<Override PartName="{part_name}" ContentType="{content_type}"/>'
    return content_types.replace("</Types>", f"{override}</Types>")


def _ensure_relationship(rels: str, rel_id: str, rel_type: str, target: str) -> str:
    if f'Type="{rel_type}"' in rels or f'Id="{rel_id}"' in rels:
        return rels
    relationship = f'<Relationship Id="{rel_id}" Type="{rel_type}" Target="{target}"/>'
    return rels.replace("</Relationships>", f"{relationship}</Relationships>")


def _notes_xml(root_name: str, item_name: str, notes: list[tuple[int, str]]) -> str:
    separator = (
        f'<w:{item_name} w:type="separator" w:id="-1">'
        "<w:p><w:r><w:separator/></w:r></w:p>"
        f"</w:{item_name}>"
    )
    continuation = (
        f'<w:{item_name} w:type="continuationSeparator" w:id="0">'
        "<w:p><w:r><w:continuationSeparator/></w:r></w:p>"
        f"</w:{item_name}>"
    )
    note_items = []
    for note_id, text in notes:
        note_items.append(
            f'<w:{item_name} w:id="{note_id}">'
            f"<w:p><w:r><w:t>{escape(text)}</w:t></w:r></w:p>"
            f"</w:{item_name}>"
        )
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f'<w:{root_name} xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f"{separator}{continuation}{''.join(note_items)}"
        f"</w:{root_name}>"
    )


def _omml_para(linear_text: str):
    escaped = (
        linear_text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )
    return parse_xml(
        f'<m:oMathPara xmlns:m="{MATH_NS}">'
        f"<m:oMath><m:r><m:t>{escaped}</m:t></m:r></m:oMath>"
        f"</m:oMathPara>"
    )
